import json
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import streamlit as st
import datetime
import os
import io
from openai import OpenAI
import PyPDF2
from typing import Optional, Dict, Any


# === FONCTIONS UTILITAIRES INTÉGRÉES ===

def valider_fichier_upload(fichier, types_autorises: list, taille_max_mb: int = 10) -> bool:
    """Valide un fichier uploadé selon les critères spécifiés"""
    if not fichier:
        return False
    
    extension = fichier.name.split('.')[-1].lower()
    if extension not in types_autorises:
        st.error(f"Type de fichier non autorisé. Extensions acceptées: {', '.join(types_autorises)}")
        return False
    
    taille_mb = len(fichier.getvalue()) / (1024 * 1024)
    if taille_mb > taille_max_mb:
        st.error(f"Fichier trop volumineux ({taille_mb:.1f} MB). Taille max: {taille_max_mb} MB")
        return False
    
    return True

def nettoyer_texte_mission(texte: str) -> str:
    """Nettoie et formate le texte de la mission pour l'analyse IA"""
    if not texte:
        return ""
    
    # Supprimer les caractères de contrôle et espaces excessifs
    texte = re.sub(r'\s+', ' ', texte)
    texte = re.sub(r'[^\w\s\-\.,:;!?\(\)\/àâäéèêëïîôöùûüÿç]', '', texte)
    texte = re.sub(r'\n\s*\n', '\n\n', texte)
    
    return texte.strip()

def detecter_domaine_mission(texte_mission: str) -> str:
    """Détecte le domaine d'activité principal ET la spécialisation de la mission"""
    domaines_keywords = {
        'Développement & Programmation': [
            'développeur', 'programmeur', 'software engineer', 'application', 'coding', 'javascript',
            'python', 'java', 'c#', 'php', 'react', 'angular', 'vue', 'node', 'backend', 'frontend',
            'fullstack', 'api', 'rest', 'mvc', 'framework', 'développement logiciel', 'programmation'
        ],
        'DevOps & Infrastructure': [
            'devops', 'infrastructure', 'cloud', 'aws', 'azure', 'gcp', 'docker', 'kubernetes',
            'ci/cd', 'jenkins', 'gitlab', 'terraform', 'ansible', 'monitoring', 'deployment',
            'orchestration', 'containerisation', 'microservices architecture'
        ],
        'Cybersécurité': [
            'cybersécurité', 'sécurité informatique', 'securité', 'pentest', 'audit sécurité',
            'security', 'firewall', 'antivirus', 'intrusion', 'vulnerability', 'iso 27001',
            'cissp', 'ethical hacking', 'forensic', 'siem', 'sox', 'gdpr compliance'
        ],
        'Intelligence Artificielle & Data': [
            'intelligence artificielle', 'machine learning', 'deep learning', 'ia', 'ai',
            'data science', 'data scientist', 'tensorflow', 'pytorch', 'scikit-learn',
            'nlp', 'computer vision', 'neural network', 'algorithme', 'big data', 'analytics'
        ],
        'Business Intelligence & Analytics': [
            'business intelligence', 'bi', 'power bi', 'tableau', 'qlik', 'reporting',
            'dashboard', 'kpi', 'data visualization', 'etl', 'data warehouse', 'olap',
            'analyse prédictive', 'data mining', 'sql server analysis services'
        ],
        'Architecture & Systèmes': [
            'architecte', 'architecture logicielle', 'système', 'enterprise architect',
            'solution architect', 'technical architect', 'patterns', 'scalabilité',
            'performance', 'haute disponibilité', 'load balancing', 'distributed systems'
        ],
        'Marketing Digital': [
            'marketing digital', 'marketing', 'communication', 'campagne', 'publicité', 'brand',
            'social media', 'seo', 'sem', 'analytics', 'crm', 'lead', 'conversion', 'content marketing',
            'google ads', 'facebook ads', 'inbound marketing'
        ],
        'Finance': [
            'finance', 'comptabilité', 'budget', 'trésorerie', 'audit', 'contrôle gestion',
            'reporting financier', 'analyse financière', 'investissement', 'risque', 'ifrs',
            'consolidation', 'fiscalité', 'treasury'
        ],
        'Ressources Humaines': [
            'ressources humaines', 'rh', 'recrutement', 'formation', 'paie', 'talent',
            'compétences', 'évaluation', 'carrière', 'mobilité', 'sirh', 'talent management',
            'people analytics', 'workforce planning'
        ],
        'Logistique & Supply Chain': [
            'logistique', 'supply chain', 'approvisionnement', 'stock', 'transport',
            'entreposage', 'distribution', 'procurement', 'planification', 'wms',
            'inventory management', 'lean', 'six sigma'
        ],
        'Consulting & Stratégie': [
            'consultant', 'conseil', 'stratégie', 'transformation', 'audit',
            'accompagnement', 'optimisation', 'expertise', 'change management',
            'business transformation', 'process improvement'
        ],
        'Santé & Médical': [
            'médical', 'santé', 'patient', 'soins', 'clinique', 'hôpital',
            'pharmacie', 'thérapie', 'diagnostic', 'healthcare', 'medical device',
            'clinical trial', 'regulatory affairs'
        ],
        'Éducation & Formation': [
            'formation', 'enseignement', 'pédagogie', 'cours', 'étudiant',
            'apprentissage', 'curriculum', 'évaluation pédagogique', 'e-learning',
            'lms', 'instructional design', 'education technology'
        ],
        'Juridique & Compliance': [
            'juridique', 'droit', 'contrat', 'compliance', 'réglementation',
            'contentieux', 'avocat', 'juriste', 'legal', 'governance',
            'risk management', 'audit compliance'
        ]
    }
    
    texte_lower = texte_mission.lower()
    scores = {}
    
    for domaine, keywords in domaines_keywords.items():
        score = sum(1 for keyword in keywords if keyword in texte_lower)
        if score > 0:
            scores[domaine] = score
    
    if scores:
        return max(scores.items(), key=lambda x: x[1])[0]
    
    return 'Général'

def extraire_categories_connaissances_par_domaine(texte_cv: str, domaine: str) -> Dict[str, list]:
    """Extrait et suggère des catégories de connaissances selon le domaine spécialisé"""
    categories_par_domaine = {
        'Développement & Programmation': {
            'Langages de programmation': ['python', 'java', 'javascript', 'c#', 'php', 'ruby', 'go', 'rust', 'kotlin'],
            'Frameworks': ['react', 'angular', 'vue', 'django', 'flask', 'spring', 'laravel', 'express'],
            'Bases de données': ['mysql', 'postgresql', 'mongodb', 'oracle', 'redis', 'elasticsearch'],
            'APIs & Services': ['rest', 'graphql', 'soap', 'microservices', 'api design', 'webhook'],
            'Outils de développement': ['git', 'github', 'vscode', 'intellij', 'postman', 'swagger'],
            'Méthodologies de dev': ['agile', 'scrum', 'kanban', 'tdd', 'bdd', 'code review']
        },
        'DevOps & Infrastructure': {
            'Outils DevOps': ['jenkins', 'gitlab ci', 'github actions', 'ansible', 'puppet', 'chef'],
            'Cloud Computing': ['aws', 'azure', 'gcp', 'serverless', 'lambda', 'cloud formation'],
            'Conteneurisation': ['docker', 'kubernetes', 'openshift', 'helm', 'docker compose'],
            'CI/CD': ['continuous integration', 'continuous deployment', 'pipeline', 'automation'],
            'Monitoring': ['prometheus', 'grafana', 'elk stack', 'nagios', 'datadog'],
            'Infrastructure as Code': ['terraform', 'cloudformation', 'arm templates', 'pulumi']
        },
        'Cybersécurité': {
            'Technologie': ['firewall', 'ids', 'ips', 'siem', 'soar', 'edr', 'antivirus', 'proxy'],
            'Normes et standards': ['iso 27001', 'iso 27002', 'nist', 'cis controls', 'pci dss', 'anssi'],
            'Réglementations': ['gdpr', 'rgpd', 'sox', 'hipaa', 'nis', 'lpm', 'dora'],
            'Outils / IDE': ['nessus', 'qualys', 'burp suite', 'metasploit', 'nmap', 'wireshark', 'kali linux'],
            'Gestion de projet': ['prince2', 'pmp', 'agile', 'scrum', 'itil', 'cobit'],
            'Langages': ['python', 'powershell', 'bash', 'sql', 'javascript', 'c++', 'java'],
            'Réseaux': ['tcp/ip', 'vpn', 'vlan', 'routing', 'switching', 'dns', 'dhcp'],
            'Systèmes d\'exploitation': ['windows', 'linux', 'unix', 'macos', 'active directory']
        },
        'Intelligence Artificielle & Data': {
            'Machine Learning': ['supervised learning', 'unsupervised learning', 'deep learning', 'nlp'],
            'Frameworks IA': ['tensorflow', 'pytorch', 'scikit-learn', 'keras', 'hugging face'],
            'Data Engineering': ['spark', 'hadoop', 'kafka', 'airflow', 'etl', 'data pipeline'],
            'Analyse de données': ['pandas', 'numpy', 'scipy', 'statistical analysis', 'data mining'],
            'Outils IA': ['jupyter', 'google colab', 'mlflow', 'kubeflow', 'azure ml'],
            'Visualisation de données': ['matplotlib', 'seaborn', 'plotly', 'tableau', 'd3.js']
        },
        'Business Intelligence & Analytics': {
            'Outils BI': ['power bi', 'tableau', 'qlik sense', 'looker', 'cognos'],
            'Reporting': ['ssrs', 'crystal reports', 'dashboard design', 'kpi monitoring'],
            'Data Warehousing': ['dimensional modeling', 'etl', 'olap', 'data mart'],
            'Visualisation': ['data visualization', 'storytelling', 'infographic', 'charts'],
            'ETL': ['ssis', 'talend', 'informatica', 'pentaho', 'data integration'],
            'Analyse prédictive': ['forecasting', 'trend analysis', 'predictive modeling']
        },
        'Architecture & Systèmes': {
            'Architecture logicielle': ['design patterns', 'solid principles', 'clean architecture'],
            'Systèmes distribués': ['microservices', 'event driven', 'cqrs', 'event sourcing'],
            'Microservices': ['api gateway', 'service mesh', 'circuit breaker', 'saga pattern'],
            'Patterns de conception': ['mvc', 'mvp', 'observer', 'factory', 'singleton'],
            'Performance': ['optimization', 'caching', 'load balancing', 'scalability'],
            'Scalabilité': ['horizontal scaling', 'vertical scaling', 'auto scaling']
        },
        'Marketing Digital': {
            'Outils marketing': ['hubspot', 'marketo', 'mailchimp', 'pardot', 'eloqua'],
            'Analytics': ['google analytics', 'adobe analytics', 'tag manager', 'heat mapping'],
            'Réseaux sociaux': ['facebook ads', 'google ads', 'linkedin ads', 'twitter ads'],
            'CRM': ['salesforce', 'hubspot', 'pipedrive', 'zoho', 'dynamics'],
            'Design graphique': ['photoshop', 'illustrator', 'canva', 'figma', 'sketch'],
            'SEO/SEM': ['seo', 'sem', 'content marketing', 'keyword research', 'link building']
        },
        'Finance': {
            'Logiciels financiers': ['sap', 'oracle financials', 'sage', 'cegid', 'blackline'],
            'Réglementation': ['ifrs', 'pcg', 'sox', 'bâle', 'mifid', 'aml'],
            'Analyse de données': ['excel avancé', 'power bi', 'tableau', 'r', 'python finance'],
            'Reporting': ['consolidation', 'business intelligence', 'financial reporting'],
            'Certification': ['cpa', 'cfa', 'frm', 'acca', 'dscg'],
            'Risk Management': ['credit risk', 'market risk', 'operational risk', 'compliance']
        },
        'Ressources Humaines': {
            'SIRH': ['sap hr', 'workday', 'adp', 'talentsoft', 'cornerstone'],
            'Recrutement': ['ats', 'linkedin recruiter', 'sourcing', 'talent acquisition'],
            'Formation': ['lms', 'e-learning', 'moodle', 'learning management', 'coaching'],
            'Paie': ['sage paie', 'adp', 'ceridian', 'payroll management'],
            'Droit social': ['droit travail', 'convention collective', 'relations sociales'],
            'Talent Management': ['performance management', 'succession planning', 'career development']
        },
        'Logistique & Supply Chain': {
            'Supply Chain': ['supply chain management', 'demand planning', 'procurement'],
            'Systèmes WMS': ['wms', 'sap wm', 'manhattan', 'warehouse management'],
            'Transport': ['tms', 'fleet management', 'logistics optimization'],
            'Réglementation': ['customs', 'trade compliance', 'incoterms'],
            'Optimisation': ['lean', 'six sigma', 'process improvement', 'inventory optimization'],
            'Procurement': ['sourcing', 'vendor management', 'contract negotiation']
        },
        'Consulting & Stratégie': {
            'Méthodologies conseil': ['mckinsey method', 'bcg matrix', 'lean startup', 'design thinking'],
            'Analyse stratégique': ['swot', 'porter five forces', 'value chain analysis'],
            'Conduite du changement': ['change management', 'kotter', 'organizational development'],
            'Gestion de projet': ['pmp', 'prince2', 'agile project management', 'scrum master'],
            'Secteurs d\'expertise': ['industry knowledge', 'domain expertise', 'market analysis'],
            'Outils d\'analyse': ['excel', 'powerpoint', 'tableau', 'power bi', 'miro']
        },
        'Santé & Médical': {
            'Systèmes médicaux': ['his', 'emr', 'ehr', 'pacs', 'ris', 'clinical systems'],
            'Réglementation santé': ['fda', 'ce marking', 'iso 13485', 'hipaa', 'gdpr santé'],
            'Dispositifs médicaux': ['medical devices', 'implants', 'diagnostic equipment'],
            'Informatique médicale': ['health informatics', 'telemedicine', 'mhealth'],
            'Qualité santé': ['gmp', 'gcp', 'quality assurance', 'risk management'],
            'Recherche clinique': ['clinical trials', 'biostatistics', 'regulatory affairs']
        },
        'Éducation & Formation': {
            'Pédagogie': ['pedagogical methods', 'learning theories', 'curriculum design'],
            'Outils e-learning': ['moodle', 'blackboard', 'canvas', 'articulate', 'captivate'],
            'Conception pédagogique': ['instructional design', 'learning objectives', 'assessment'],
            'Évaluation': ['evaluation methods', 'learning analytics', 'competency assessment'],
            'Technologies éducatives': ['edtech', 'virtual classroom', 'gamification'],
            'Ingénierie de formation': ['training needs analysis', 'learning path design']
        },
        'Juridique & Compliance': {
            'Droit des affaires': ['corporate law', 'contract law', 'commercial law'],
            'Compliance': ['regulatory compliance', 'internal controls', 'policy development'],
            'Réglementation': ['gdpr', 'sox', 'anti money laundering', 'trade sanctions'],
            'Contentieux': ['litigation', 'dispute resolution', 'arbitration'],
            'Propriété intellectuelle': ['intellectual property', 'patents', 'trademarks'],
            'Outils juridiques': ['legal research', 'case management', 'contract management']
        }
    }
    
    if domaine not in categories_par_domaine:
        return {
            'Outils et Logiciels': [],
            'Compétences Techniques': [],
            'Méthodologies': []
        }
    
    texte_lower = texte_cv.lower()
    categories_trouvees = {}
    
    for categorie, keywords in categories_par_domaine[domaine].items():
        mots_trouves = [keyword for keyword in keywords if keyword in texte_lower]
        if mots_trouves:
            categories_trouvees[categorie] = mots_trouves
    
    return categories_trouvees

def calculer_score_adequation(dossier_competences: str, texte_mission: str) -> float:
    """Calcule un score d'adéquation entre le dossier et la mission"""
    domaine = detecter_domaine_mission(texte_mission)
    categories = extraire_categories_connaissances_par_domaine(dossier_competences, domaine)
    
    score_total = sum(len(mots) for mots in categories.values() if mots)
    score_max = sum(len(keywords) for keywords in categories.values())
    
    return min(score_total / max(score_max, 1), 1.0) if score_max > 0 else 0

def generer_rapport_optimisation(donnees_originales: str, donnees_optimisees: dict, 
                                 texte_mission: str) -> dict:
    """Génère un rapport d'optimisation détaillé"""
    domaine = detecter_domaine_mission(texte_mission)
    categories = extraire_categories_connaissances_par_domaine(donnees_originales, domaine)
    
    rapport = {
        'domaine_detecte': domaine,
        'score_adequation': calculer_score_adequation(donnees_originales, texte_mission),
        'categories_identifiees': list(categories.keys()),
        'nb_experiences': len(donnees_optimisees.get('experiences', [])),
        'nb_formations': len(donnees_optimisees.get('formations', [])),
        'nb_points_forts': len(donnees_optimisees.get('points_forts', [])),
        'categories_competences': list(donnees_optimisees.get('connaissances', {}).keys())
    }
    
    return rapport

def sauvegarder_historique_generation(donnees: dict, nom_fichier: str, rapport: dict):
    """Sauvegarde l'historique des générations"""
    try:
        historique = {
            'timestamp': datetime.datetime.now().isoformat(),
            'nom_consultant': donnees.get('nom_consultant', 'Inconnu'),
            'titre_poste': donnees.get('titre_du_poste', 'Inconnu'),
            'domaine_mission': rapport.get('domaine_detecte', 'Non défini'),
            'score_adequation': rapport.get('score_adequation', 0),
            'nom_fichier': nom_fichier,
            'nb_experiences': len(donnees.get('experiences', [])),
            'nb_formations': len(donnees.get('formations', []))
        }
        
        st.info(f"📝 Génération enregistrée : {historique['nom_consultant']} - {historique['domaine_mission']} - Score: {historique['score_adequation']:.1%}")
        
    except Exception as e:
        st.warning(f"Erreur lors de la sauvegarde : {str(e)}")

# === CONFIGURATION OPENAI ===

def configurer_openai():
    try:
        import openai
        st.write(f"Version OpenAI: {openai.__version__}")
        
        from openai import OpenAI
        api_key = st.secrets.get("OPENAI_API_KEY")
        
        # Essai minimal
        client = OpenAI(api_key=api_key)
        return client
        
    except Exception as e:
        st.write(f"Erreur détaillée: {str(e)}")
        st.write(f"Type d'erreur: {type(e)}")
        return None
    
# === LECTURE DE FICHIERS ===

def lire_fichier_pdf(fichier_pdf) -> str:
    """Extrait le texte d'un fichier PDF"""
    try:
        pdf_reader = PyPDF2.PdfReader(fichier_pdf)
        texte = ""
        for page in pdf_reader.pages:
            texte += page.extract_text() + "\n"
        return texte.strip()
    except Exception as e:
        st.error(f"Erreur lors de la lecture du PDF : {str(e)}")
        return ""

def lire_fichier_txt(fichier_txt) -> str:
    """Lit le contenu d'un fichier texte"""
    try:
        encodings = ['utf-8', 'latin-1', 'cp1252']
        for encoding in encodings:
            try:
                fichier_txt.seek(0)
                contenu = fichier_txt.read()
                if isinstance(contenu, bytes):
                    return contenu.decode(encoding)
                return contenu
            except UnicodeDecodeError:
                continue
        return ""
    except Exception as e:
        st.error(f"Erreur lors de la lecture du fichier texte : {str(e)}")
        return ""

def lire_fichier_word(fichier_word) -> str:
    """Extrait le texte d'un fichier Word"""
    try:
        doc = Document(fichier_word)
        texte = ""
        
        # Extraire le texte des paragraphes
        for paragraph in doc.paragraphs:
            texte += paragraph.text + "\n"
        
        # Extraire le texte des tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texte += cell.text + " "
                texte += "\n"
        
        return texte.strip()
    except Exception as e:
        st.error(f"Erreur lors de la lecture du fichier Word : {str(e)}")
        return ""

# === GÉNÉRATION DE PROMPTS ET APPELS OPENAI ===

def generer_prompt_optimisation(description_mission: str, dossier_competences: str) -> str:
    """Génère le prompt pour optimiser le dossier de compétences selon la mission"""
    
    prompt = f"""
Tu es un expert RH spécialisé dans l'optimisation de dossiers de compétences pour des missions spécifiques.

**MISSION À ANALYSER :**
{description_mission}

**DOSSIER DE COMPÉTENCES ACTUEL :**
{dossier_competences}

**OBJECTIF :**
Adapter et optimiser le dossier de compétences pour qu'il soit parfaitement aligné avec la mission décrite, en mettant en valeur les compétences, expériences et qualités les plus pertinentes.

**INSTRUCTIONS CRITIQUES :**
1. **ANALYSER** la mission en profondeur pour identifier les compétences clés requises
2. **CONSERVER** toutes les informations authentiques du dossier original
3. **REFORMULER** et **PRIORISER** les éléments les plus pertinents pour la mission
4. **ENRICHIR** les descriptions pour montrer l'adéquation avec les besoins
5. **OPTIMISER** chaque section pour maximiser l'impact
6. **GÉNÉRER** automatiquement les catégories de connaissances selon le domaine détecté

**INSTRUCTIONS SPÉCIALES POUR LES CONNAISSANCES :**
- ANALYSER le domaine de la mission et sa spécialisation exacte
- CRÉER des catégories de compétences pertinentes pour cette spécialisation spécifique
- Les catégories doivent être adaptées au secteur d'activité ET à la spécialisation identifiée
- Utiliser les compétences réelles du dossier pour peupler ces catégories

**EXEMPLES DE CATÉGORIES SELON LES DOMAINES ET SPÉCIALISATIONS :**

**DÉVELOPPEMENT & PROGRAMMATION :**
- "Langages de programmation", "Frameworks", "Bases de données", "APIs & Services", "Outils de développement", "Méthodologies de dev"

**DEVOPS & INFRASTRUCTURE :**
- "Outils DevOps", "Cloud Computing", "Conteneurisation", "CI/CD", "Monitoring", "Infrastructure as Code"

**CYBERSÉCURITÉ :**
- "Sécurité réseau", "Outils de sécurité", "Audit & Conformité", "Gestion des risques", "Certifications sécurité", "Tests d'intrusion"

**INTELLIGENCE ARTIFICIELLE & DATA :**
- "Machine Learning", "Frameworks IA", "Data Engineering", "Analyse de données", "Outils IA", "Visualisation de données"

**BUSINESS INTELLIGENCE & ANALYTICS :**
- "Outils BI", "Reporting", "Data Warehousing", "Visualisation", "ETL", "Analyse prédictive"

**ARCHITECTURE & SYSTÈMES :**
- "Architecture logicielle", "Systèmes distribués", "Microservices", "Patterns de conception", "Performance", "Scalabilité"

**MARKETING DIGITAL :**
- "Outils marketing", "Analytics", "Réseaux sociaux", "CRM", "Design graphique", "SEO/SEM"

**FINANCE :**
- "Logiciels financiers", "Réglementation", "Analyse de données", "Reporting", "Certification", "Risk Management"

**RESSOURCES HUMAINES :**
- "SIRH", "Recrutement", "Formation", "Paie", "Droit social", "Talent Management"

**LOGISTIQUE & SUPPLY CHAIN :**
- "Supply Chain", "Systèmes WMS", "Transport", "Réglementation", "Optimisation", "Procurement"

**CONSULTING & STRATÉGIE :**
- "Méthodologies conseil", "Analyse stratégique", "Conduite du changement", "Gestion de projet", "Secteurs d'expertise", "Outils d'analyse"

**SANTÉ & MÉDICAL :**
- "Systèmes médicaux", "Réglementation santé", "Dispositifs médicaux", "Informatique médicale", "Qualité santé", "Recherche clinique"

**ÉDUCATION & FORMATION :**
- "Pédagogie", "Outils e-learning", "Conception pédagogique", "Évaluation", "Technologies éducatives", "Ingénierie de formation"

**JURIDIQUE & COMPLIANCE :**
- "Droit des affaires", "Compliance", "Réglementation", "Contentieux", "Propriété intellectuelle", "Outils juridiques"

**FORMAT DE SORTIE OBLIGATOIRE - JSON STRICT :**
{{
  "nom_consultant": "Nom complet du consultant",
  "titre_du_poste": "Titre optimisé pour la mission",
  "points_forts": [
    "Point fort 1 adapté à la mission",
    "Point fort 2 adapté à la mission",
    "Point fort 3 adapté à la mission",
    "Point fort 4 adapté à la mission",
    "Point fort 5 adapté à la mission",
    "Point fort 6 adapté à la mission"
  ],
  "niveaux_intervention": [
    "Niveau 1 correspondant à la mission",
    "Niveau 2 correspondant à la mission",
    "Niveau 3 correspondant à la mission",
    "Niveau 4 correspondant à la mission",
    "Niveau 5 correspondant à la mission"
  ],
  "formations": [
    {{
      "annee": "YYYY",
      "intitule": "Formation pertinente pour la mission"
    }}
  ],
  "connaissances": {{
    "Catégorie 1 adaptée à la spécialisation": "Compétences spécifiques extraites du dossier",
    "Catégorie 2 adaptée à la spécialisation": "Compétences spécifiques extraites du dossier",
    "Catégorie 3 adaptée à la spécialisation": "Compétences spécifiques extraites du dossier",
    "Catégorie 4 adaptée à la spécialisation": "Compétences spécifiques extraites du dossier",
    "Catégorie 5 adaptée à la spécialisation": "Compétences spécifiques extraites du dossier"
  }},
  "hobbies_divers": {{
    "langues": "Langues maîtrisées",
    "hobbies": "Hobbies valorisant le profil"
  }},
  "experiences": [
    {{
      "periode": "Période de l'expérience",
      "titre": "Titre optimisé pour la mission",
      "entreprise": "Nom de l'entreprise",
      "responsabilites": "Responsabilités adaptées et détaillées pour la mission",
      "realisations": [
        "Réalisation 1 alignée avec la mission",
        "Réalisation 2 alignée avec la mission",
        "Réalisation 3 alignée avec la mission",
        "Réalisation 4 alignée avec la mission",
        "Réalisation 5 alignée avec la mission",
        "Réalisation 6 alignée avec la mission",
        "Réalisation 7 alignée avec la mission",
        "Réalisation 8 alignée avec la mission",
        "Réalisation 9 alignée avec la mission",
        "Réalisation 10 alignée avec la mission"
      ],
      "environnement": "Environnement technique pertinent"
    }}
  ]
}}

**RÈGLES IMPORTANTES :**
- Utilise UNIQUEMENT les informations réelles du dossier original
- Les catégories de connaissances DOIVENT être générées selon la spécialisation ou les spécialisations exactes identifiées
- Reformule intelligemment pour mettre en valeur l'adéquation
- Priorise les éléments les plus pertinents pour la mission
- Assure-toi que le JSON est parfaitement formaté
- N'invente AUCUNE information, adapte seulement ce qui existe
- Génère le maximum de catégories de connaissances pertinentes pour la spécialisation
- Sois PRÉCIS dans le choix des catégories selon la spécialisation détectée
- **RÈGLE CRITIQUE D'ALIGNEMENT** : Chaque expérience, point fort et niveau d'intervention inclus DOIT avoir un lien évident et démontrable avec les besoins exprimés dans la mission
- Si un élément du CV n'apporte aucune valeur pour cette mission spécifique : l'EXCLURE totalement
- Génère le maximum de 10 réalisations par experiences 
- Adapter les catégories de compétences si le profil couvre plusieurs domaines (ex: DevSecOps = DevOps + Sécurité)
"""
    
    return prompt

def appeler_openai_pour_optimisation(description_mission: str, dossier_competences: str) -> Optional[Dict[Any, Any]]:
    """Appelle l'API OpenAI pour optimiser le dossier de compétences"""
    
    client = configurer_openai()
    if not client:
        return None
    
    try:
        prompt = generer_prompt_optimisation(description_mission, dossier_competences)
        
        with st.spinner("🤖 Analyse intelligente en cours avec OpenAI..."):
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system", 
                        "content": "Tu es un expert RH spécialisé dans l'optimisation de dossiers de compétences. Réponds UNIQUEMENT avec un JSON valide, sans texte supplémentaire."
                    },
                    {
                        "role": "user", 
                        "content": prompt
                    }
                ],
                temperature=0.3,
                max_tokens=4000
            )
            
            contenu_reponse = response.choices[0].message.content.strip()
            
            # Nettoyer la réponse pour extraire le JSON
            if contenu_reponse.startswith("```json"):
                contenu_reponse = contenu_reponse[7:-3].strip()
            elif contenu_reponse.startswith("```"):
                contenu_reponse = contenu_reponse[3:-3].strip()
            
            # Parser le JSON
            try:
                donnees_optimisees = json.loads(contenu_reponse)
                return donnees_optimisees
            except json.JSONDecodeError as e:
                st.error(f"Erreur de parsing JSON : {e}")
                st.text("Réponse brute reçue :")
                st.text(contenu_reponse[:500] + "..." if len(contenu_reponse) > 500 else contenu_reponse)
                return None
                
    except Exception as e:
        st.error(f"Erreur lors de l'appel à OpenAI : {str(e)}")
        return None
# === EXTRACTION DE DONNÉES JSON ===

def extraire_contenu_json(reponse_ia):
    """Extrait le contenu JSON de la réponse de l'IA - Version améliorée"""
    # Si c'est déjà un dictionnaire, le retourner directement
    if isinstance(reponse_ia, dict):
        return reponse_ia
    
    # Si c'est une chaîne JSON directe
    if isinstance(reponse_ia, str):
        try:
            return json.loads(reponse_ia)
        except json.JSONDecodeError:
            pass
    
    # Patterns pour extraire le JSON de différents formats de réponse
    patterns = [
        r'content=\'({.*?})\'',
        r'```json\s*({.*?})\s*```',
        r'```\s*({.*?})\s*```',
        r'({[^{}]*"nom_consultant"[^}]*})'
    ]
    
    for pattern in patterns:
        match = re.search(pattern, str(reponse_ia), re.DOTALL)
        if match:
            json_str = match.group(1)
            json_str = json_str.replace('\\\\', '\\')
            json_str = json_str.replace("\\'", "'")
            json_str = json_str.replace('\\n', '\n')
            json_str = re.sub(r'\\(?!["\\/bfnrt])', r'\\\\', json_str)
            
            try:
                return json.loads(json_str)
            except json.JSONDecodeError:
                continue
    
    # Tentative d'extraction manuelle en dernier recours
    return extraire_donnees_manuellement(str(reponse_ia))

def extraire_donnees_manuellement(json_str):
    """Extraction manuelle des données en cas d'échec du parsing JSON"""
    donnees = {}
    champs = [
        'nom_consultant', 'titre_du_poste', 'niveaux_intervention',
        'hobbies_et_divers', 'connaissances', 'experiences', 
        'mois_debut_experience', 'nom_entreprise', 'points_forts', 'formations'
    ]
    
    for champ in champs:
        pattern = rf'"{champ}":\s*"([^"]*(?:\\.[^"]*)*)"'
        match = re.search(pattern, json_str, re.DOTALL)
        if match:
            valeur = match.group(1)
            valeur = valeur.replace('\\"', '"').replace('\\n', '\n').replace("\\'", "'")
            donnees[champ] = valeur
    
    return donnees if donnees else None

# Toutes les autres fonctions du fichier original (remplacer_texte_paragraph, ajouter_liste_paragraphes, 
# creer_tableau_*, generer_cv_depuis_template_avec_entete_preserve, etc.) restent identiques...

# [Ici vous pouvez ajouter toutes les autres fonctions de génération Word de votre fichier original]

def extraire_contenu_json(reponse_ia):
    """
    Extrait le contenu JSON de la réponse de l'IA
    Version améliorée pour gérer les réponses directes d'OpenAI
    """
    # Si c'est déjà un dictionnaire, le retourner directement
    if isinstance(reponse_ia, dict):
        return reponse_ia
    
    # Si c'est une chaîne JSON directe
    if isinstance(reponse_ia, str):
        try:
            return json.loads(reponse_ia)
        except json.JSONDecodeError:
            pass
    
    # Patterns pour extraire le JSON de différents formats de réponse
    patterns = [
        r'content=\'({.*?})\'',
        r'```json\s*({.*?})\s*```',
        r'```\s*({.*?})\s*```',
        r'({[^{}]*"nom_consultant"[^}]*})'
    ]
    
    for pattern in patterns:
        match = re.search(pattern, str(reponse_ia), re.DOTALL)
        if match:
            json_str = match.group(1)
            json_str = json_str.replace('\\\\', '\\')
            json_str = json_str.replace("\\'", "'")
            json_str = json_str.replace('\\n', '\n')
            json_str = re.sub(r'\\(?!["\\/bfnrt])', r'\\\\', json_str)
            
            try:
                return json.loads(json_str)
            except json.JSONDecodeError as e:
                st.error(f"Erreur de parsing JSON: {e}")
                continue
    
    # Tentative d'extraction manuelle en dernier recours
    return extraire_donnees_manuellement(str(reponse_ia))

def extraire_donnees_manuellement(json_str):
    """
    Extraction manuelle des données en cas d'échec du parsing JSON
    """
    donnees = {}
    champs = [
        'nom_consultant', 'titre_du_poste', 'niveaux_intervention',
        'hobbies_et_divers', 'connaissances', 'experiences', 
        'mois_debut_experience', 'nom_entreprise', 'points_forts', 'formations'
    ]
    
    for champ in champs:
        pattern = rf'"{champ}":\s*"([^"]*(?:\\.[^"]*)*)"'
        match = re.search(pattern, json_str, re.DOTALL)
        if match:
            valeur = match.group(1)
            valeur = valeur.replace('\\"', '"').replace('\\n', '\n').replace("\\'", "'")
            donnees[champ] = valeur
    
    return donnees if donnees else None

# [Conserver toutes les autres fonctions du fichier original ici...]
# remplacer_texte_paragraph, ajouter_liste_paragraphes, creer_tableau_*, etc.

def extraire_contenu_json(reponse_ia):
    """
    Extrait le contenu JSON de la réponse de l'IA
    """
    pattern = r'content=\'({.*?})\''
    match = re.search(pattern, str(reponse_ia), re.DOTALL)
    
    if match:
        json_str = match.group(1)
        json_str = json_str.replace('\\\\', '\\')
        json_str = json_str.replace("\\'", "'")
        json_str = json_str.replace('\\n', '\n')
        json_str = re.sub(r'\\(?!["\\/bfnrt])', r'\\\\', json_str)
        
        try:
            return json.loads(json_str)
        except json.JSONDecodeError as e:
            st.error(f"Erreur de parsing JSON: {e}")
            try:
                json_str_corrige = json_str.replace("\\", "\\\\")
                json_str_corrige = json_str_corrige.replace('\\"', '"')
                json_str_corrige = json_str_corrige.replace("\\\\n", "\\n")
                json_str_corrige = json_str_corrige.replace("\\\\'", "'")
                return json.loads(json_str_corrige)
            except json.JSONDecodeError:
                st.warning("Tentative d'extraction manuelle des données...")
                return extraire_donnees_manuellement(json_str)
    else:
        try:
            pattern2 = r'({[^{}]*"nom_consultant"[^}]*})'
            match2 = re.search(pattern2, str(reponse_ia), re.DOTALL)
            if match2:
                return json.loads(match2.group(1))
        except:
            pass
        
        st.error("Aucun contenu JSON trouvé dans la réponse")
        return None

def extraire_donnees_manuellement(json_str):
    """
    Extraction manuelle des données en cas d'échec du parsing JSON
    """
    donnees = {}
    champs = [
        'nom_consultant', 'titre_du_poste', 'niveaux_intervention',
        'hobbies_et_divers', 'connaissances', 'experiences', 
        'mois_debut_experience', 'nom_entreprise', 'points_forts', 'formations'
    ]
    
    for champ in champs:
        pattern = rf'"{champ}":\s*"([^"]*(?:\\.[^"]*)*)"'
        match = re.search(pattern, json_str, re.DOTALL)
        if match:
            valeur = match.group(1)
            valeur = valeur.replace('\\"', '"').replace('\\n', '\n').replace("\\'", "'")
            donnees[champ] = valeur
    
    return donnees if donnees else None

def remplacer_texte_paragraph(paragraph, ancien_texte, nouveau_texte):
    """
    Remplace du texte dans un paragraphe en préservant le formatage
    """
    if ancien_texte in paragraph.text:
        for run in paragraph.runs:
            if ancien_texte in run.text:
                run.text = run.text.replace(ancien_texte, nouveau_texte)
                return True
        
        texte_complet = paragraph.text
        if ancien_texte in texte_complet:
            paragraph.clear()
            nouveau_run = paragraph.add_run(texte_complet.replace(ancien_texte, nouveau_texte))
            return True
    
    return False

def ajouter_liste_paragraphes(doc, paragraph_parent, items, style_bullet=True):
    """
    Ajoute une série de paragraphes formatés en liste avec puces en gras
    et texte en format normal
    Sans toucher au paragraphe parent existant
    """
    if not items:
        return True
    
    parent_element = paragraph_parent._element.getparent()
    insert_index = list(parent_element).index(paragraph_parent._element)
    couleur_clinkast = "1F4E79"
    
    # Modifier aussi le paragraphe parent pour qu'il ait les mêmes propriétés
    parent_p_element = paragraph_parent._element
    
    # Modifier le formatage du contenu existant du paragraphe parent
    for run in parent_p_element.findall(qn('w:r')):
        r_props = run.find(qn('w:rPr'))
        if r_props is not None:
            # Supprimer le gras existant des runs de texte
            for child in list(r_props):
                if child.tag == qn('w:b'):
                    r_props.remove(child)
        
        # Traiter le texte pour séparer puce et contenu si nécessaire
        text_elem = run.find(qn('w:t'))
        if text_elem is not None and text_elem.text and text_elem.text.startswith('•'):
            # Diviser en deux runs : puce en gras + texte normal
            original_text = text_elem.text
            if len(original_text) > 3:  # "•\u00A0\u00A0" + contenu
                # Modifier le run existant pour ne contenir que la puce avec espaces insécables
                text_elem.text = "•\u00A0\u00A0"  # Puce + 2 espaces insécables
                
                # Recréer complètement les propriétés du run de la puce
                if r_props is None:
                    r_props = OxmlElement('w:rPr')
                    run.insert(0, r_props)
                else:
                    # Nettoyer les propriétés existantes
                    for child in list(r_props):
                        r_props.remove(child)
                
                # Appliquer toutes les propriétés uniformes pour la puce
                r_bullet_font = OxmlElement('w:rFonts')
                r_bullet_font.set(qn('w:ascii'), 'Calibri')
                r_bullet_font.set(qn('w:hAnsi'), 'Calibri')
                r_bullet_bold = OxmlElement('w:b')
                r_bullet_color = OxmlElement('w:color')
                r_bullet_color.set(qn('w:val'), couleur_clinkast)
                r_bullet_size = OxmlElement('w:sz')
                r_bullet_size.set(qn('w:val'), '20')  # 10pt
                
                r_props.append(r_bullet_font)
                r_props.append(r_bullet_bold)
                r_props.append(r_bullet_color)
                r_props.append(r_bullet_size)
                
                # Créer un nouveau run pour le texte
                new_text_run = OxmlElement('w:r')
                new_text_props = OxmlElement('w:rPr')
                
                # Propriétés du nouveau run (sans gras)
                r_text_color = OxmlElement('w:color')
                r_text_color.set(qn('w:val'), couleur_clinkast)
                r_text_size = OxmlElement('w:sz')
                r_text_size.set(qn('w:val'), '20')
                r_text_font = OxmlElement('w:rFonts')
                r_text_font.set(qn('w:ascii'), 'Calibri')
                r_text_font.set(qn('w:hAnsi'), 'Calibri')
                
                new_text_props.append(r_text_font)
                new_text_props.append(r_text_color)
                new_text_props.append(r_text_size)
                new_text_run.append(new_text_props)
                
                # Texte sans la puce
                new_text_elem = OxmlElement('w:t')
                new_text_elem.text = original_text[2:]  # Enlever "• "
                new_text_run.append(new_text_elem)
                
                # Insérer le nouveau run après le run actuel
                parent_index = list(parent_p_element).index(run)
                parent_p_element.insert(parent_index + 1, new_text_run)
    
    # Vérifier si le paragraphe parent a des propriétés, sinon les créer
    parent_p_props = parent_p_element.find(qn('w:pPr'))
    if parent_p_props is None:
        parent_p_props = OxmlElement('w:pPr')
        parent_p_element.insert(0, parent_p_props)
    
    # Nettoyer les anciennes propriétés d'indentation et espacement
    for child in list(parent_p_props):
        if child.tag in [qn('w:ind'), qn('w:spacing'), qn('w:jc'), qn('w:pStyle'), qn('w:pBdr')]:
            parent_p_props.remove(child)
    
    # Appliquer les mêmes propriétés au paragraphe parent
    parent_p_align = OxmlElement('w:jc')
    parent_p_align.set(qn('w:val'), 'left')
    parent_p_props.append(parent_p_align)
    
    parent_p_ind = OxmlElement('w:ind')
    parent_p_ind.set(qn('w:left'), '0')
    parent_p_ind.set(qn('w:leftChars'), '0')
    parent_p_ind.set(qn('w:firstLine'), '0')
    parent_p_ind.set(qn('w:firstLineChars'), '0')
    parent_p_ind.set(qn('w:hanging'), '0')
    parent_p_ind.set(qn('w:hangingChars'), '0')
    parent_p_ind.set(qn('w:right'), '0')
    parent_p_ind.set(qn('w:rightChars'), '0')
    parent_p_props.append(parent_p_ind)
    
    parent_p_spacing = OxmlElement('w:spacing')
    parent_p_spacing.set(qn('w:before'), '0')
    parent_p_spacing.set(qn('w:beforeLines'), '0')
    parent_p_spacing.set(qn('w:after'), '0')
    parent_p_spacing.set(qn('w:afterLines'), '0')
    parent_p_spacing.set(qn('w:line'), '240')
    parent_p_spacing.set(qn('w:lineRule'), 'auto')
    parent_p_props.append(parent_p_spacing)
    
    parent_p_pbdr = OxmlElement('w:pBdr')
    parent_p_props.append(parent_p_pbdr)
    
    parent_p_style = OxmlElement('w:pStyle')
    parent_p_style.set(qn('w:val'), 'Normal')
    parent_p_props.append(parent_p_style)
    
    for i, item in enumerate(items):
        new_p_element = parent_element.makeelement(qn('w:p'))
        
        # Propriétés du paragraphe
        p_props = OxmlElement('w:pPr')
        
        # Alignement à gauche
        p_align = OxmlElement('w:jc')
        p_align.set(qn('w:val'), 'left')
        p_props.append(p_align)
        
        # Supprimer complètement toute indentation
        p_ind = OxmlElement('w:ind')
        p_ind.set(qn('w:left'), '0')
        p_ind.set(qn('w:leftChars'), '0')
        p_ind.set(qn('w:firstLine'), '0')
        p_ind.set(qn('w:firstLineChars'), '0')
        p_ind.set(qn('w:hanging'), '0')
        p_ind.set(qn('w:hangingChars'), '0')
        p_ind.set(qn('w:right'), '0')
        p_ind.set(qn('w:rightChars'), '0')
        p_props.append(p_ind)
        
        # Supprimer l'espacement avant et après
        p_spacing = OxmlElement('w:spacing')
        p_spacing.set(qn('w:before'), '0')
        p_spacing.set(qn('w:beforeLines'), '0')
        p_spacing.set(qn('w:after'), '0')
        p_spacing.set(qn('w:afterLines'), '0')
        p_spacing.set(qn('w:line'), '240')
        p_spacing.set(qn('w:lineRule'), 'auto')
        p_props.append(p_spacing)
        
        # Ajouter des marges de paragraphe nulles
        p_pbdr = OxmlElement('w:pBdr')
        p_props.append(p_pbdr)
        
        # Forcer le style à Normal pour éviter l'héritage
        p_style = OxmlElement('w:pStyle')
        p_style.set(qn('w:val'), 'Normal')
        p_props.append(p_style)
        
        new_p_element.append(p_props)
        
        if style_bullet:
            # Premier run : la puce en gras
            r_bullet_element = OxmlElement('w:r')
            r_bullet_props = OxmlElement('w:rPr')
            
            # Propriétés du run pour la puce (en gras)
            r_bullet_bold = OxmlElement('w:b')
            r_bullet_color = OxmlElement('w:color')
            r_bullet_color.set(qn('w:val'), couleur_clinkast)
            r_bullet_size = OxmlElement('w:sz')
            r_bullet_size.set(qn('w:val'), '20')  # 10pt
            r_bullet_font = OxmlElement('w:rFonts')
            r_bullet_font.set(qn('w:ascii'), 'Calibri')
            r_bullet_font.set(qn('w:hAnsi'), 'Calibri')
            
            r_bullet_props.append(r_bullet_font)
            r_bullet_props.append(r_bullet_bold)
            r_bullet_props.append(r_bullet_color)
            r_bullet_props.append(r_bullet_size)
            r_bullet_element.append(r_bullet_props)
            
            # Texte de la puce uniquement
            r_bullet_text = OxmlElement('w:t')
            r_bullet_text.text = "•\u00A0\u00A0"
            r_bullet_element.append(r_bullet_text)
            
            new_p_element.append(r_bullet_element)
            
            # Deuxième run : le texte en format normal
            r_text_element = OxmlElement('w:r')
            r_text_props = OxmlElement('w:rPr')
            
            # Propriétés du run pour le texte (format normal, sans gras)
            r_text_color = OxmlElement('w:color')
            r_text_color.set(qn('w:val'), couleur_clinkast)
            r_text_size = OxmlElement('w:sz')
            r_text_size.set(qn('w:val'), '20')  # 10pt
            r_text_font = OxmlElement('w:rFonts')
            r_text_font.set(qn('w:ascii'), 'Calibri')
            r_text_font.set(qn('w:hAnsi'), 'Calibri')
            
            r_text_props.append(r_text_font)
            r_text_props.append(r_text_color)
            r_text_props.append(r_text_size)
            # Pas de r_bold ici pour le texte
            r_text_element.append(r_text_props)
            
            # Texte de l'item
            r_text_text = OxmlElement('w:t')
            r_text_text.text = item.strip()
            r_text_element.append(r_text_text)
            
            new_p_element.append(r_text_element)
        else:
            # Run unique sans puce, texte normal
            r_element = OxmlElement('w:r')
            r_props = OxmlElement('w:rPr')
            
            # Propriétés du run (format normal, sans gras)
            r_color = OxmlElement('w:color')
            r_color.set(qn('w:val'), couleur_clinkast)
            r_size = OxmlElement('w:sz')
            r_size.set(qn('w:val'), '20')  # 10pt
            r_font = OxmlElement('w:rFonts')
            r_font.set(qn('w:ascii'), 'Calibri')
            r_font.set(qn('w:hAnsi'), 'Calibri')
            
            r_props.append(r_font)
            r_props.append(r_color)
            r_props.append(r_size)
            # Pas de r_bold ici
            r_element.append(r_props)
            
            # Texte
            r_text = OxmlElement('w:t')
            r_text.text = item.strip()
            r_element.append(r_text)
            
            new_p_element.append(r_element)
        
        # Insérer le paragraphe
        parent_element.insert(insert_index + 1 + i, new_p_element)
    
    return True

def creer_tableau_connaissances_a_position(doc, paragraph_position, connaissances_dict):
    """
    Crée un tableau de connaissances à une position spécifique dans le document
    """
    if not connaissances_dict or not isinstance(connaissances_dict, dict):
        connaissances_dict = {
            'Langages et Framework': '.NET (C#, ASP.NET), MVC, WEB API, ANGULAR, TYPESCRIPT',
            'SGBD': 'MYSQL, POSTGRESQL, MONGODB, SQL Serveur',
            'Systèmes d\'exploitation': 'Linux (Ubuntu), Windows',
            'Outils': 'VsCode, GIT, GitHub, Visual studio',
            'DevOps et Cloud': 'DOCKER, KUBERNETES, CI/CD Devops',
            'Méthodologie': 'Agile SCRUM'
        }
    
    parent_element = paragraph_position._element.getparent()
    paragraph_index = list(parent_element).index(paragraph_position._element)
    
    table_element = parent_element.makeelement(qn('w:tbl'))
    
    tbl_props = OxmlElement('w:tblPr')
    tbl_style = OxmlElement('w:tblStyle')
    tbl_style.set(qn('w:val'), 'TableGrid')
    tbl_props.append(tbl_style)
    
    tbl_width = OxmlElement('w:tblW')
    tbl_width.set(qn('w:w'), '5000')
    tbl_width.set(qn('w:type'), 'pct')
    tbl_props.append(tbl_width)
    
    table_element.append(tbl_props)
    
    couleur_clinkast = "1F4E79"
    
    for categorie, contenu in connaissances_dict.items():
        tr_element = OxmlElement('w:tr')
        
        # Cellule 1 - Catégorie
        tc1_element = OxmlElement('w:tc')
        tc1_props = OxmlElement('w:tcPr')
        
        tc1_width = OxmlElement('w:tcW')
        tc1_width.set(qn('w:w'), '2000')
        tc1_width.set(qn('w:type'), 'pct')
        tc1_props.append(tc1_width)
        
        # Supprimer complètement les marges internes
        tc1_margins = OxmlElement('w:tcMar')
        tc1_left = OxmlElement('w:left')
        tc1_left.set(qn('w:w'), '0')
        tc1_left.set(qn('w:type'), 'dxa')
        tc1_right = OxmlElement('w:right')
        tc1_right.set(qn('w:w'), '0')
        tc1_right.set(qn('w:type'), 'dxa')
        tc1_top = OxmlElement('w:top')
        tc1_top.set(qn('w:w'), '0')
        tc1_top.set(qn('w:type'), 'dxa')
        tc1_bottom = OxmlElement('w:bottom')
        tc1_bottom.set(qn('w:w'), '120')
        tc1_bottom.set(qn('w:type'), 'dxa')
        tc1_margins.append(tc1_left)
        tc1_margins.append(tc1_right)
        tc1_margins.append(tc1_top)
        tc1_margins.append(tc1_bottom)
        tc1_props.append(tc1_margins)
        
        tc1_element.append(tc1_props)
        
        p1_element = OxmlElement('w:p')
        p1_props = OxmlElement('w:pPr')
        p1_align = OxmlElement('w:jc')
        p1_align.set(qn('w:val'), 'left')
        p1_props.append(p1_align)
        
        # Supprimer complètement toute indentation et espacement (connaissances - plus strict)
        p1_ind = OxmlElement('w:ind')
        p1_ind.set(qn('w:left'), '0')
        p1_ind.set(qn('w:leftChars'), '0')
        p1_ind.set(qn('w:firstLine'), '0')
        p1_ind.set(qn('w:firstLineChars'), '0')
        p1_ind.set(qn('w:hanging'), '0')
        p1_ind.set(qn('w:hangingChars'), '0')
        p1_ind.set(qn('w:right'), '0')
        p1_ind.set(qn('w:rightChars'), '0')
        p1_props.append(p1_ind)
        
        # Supprimer l'espacement avant et après (plus strict)
        p1_spacing = OxmlElement('w:spacing')
        p1_spacing.set(qn('w:before'), '0')
        p1_spacing.set(qn('w:beforeLines'), '0')
        p1_spacing.set(qn('w:after'), '0')
        p1_spacing.set(qn('w:afterLines'), '0')
        p1_spacing.set(qn('w:line'), '240')
        p1_spacing.set(qn('w:lineRule'), 'auto')
        p1_props.append(p1_spacing)
        
        # Ajouter des marges de paragraphe nulles
        p1_pbdr = OxmlElement('w:pBdr')
        p1_props.append(p1_pbdr)
        
        p1_element.append(p1_props)
        
        r1_element = OxmlElement('w:r')
        r1_props = OxmlElement('w:rPr')
        r1_bold = OxmlElement('w:b')
        r1_color = OxmlElement('w:color')
        r1_color.set(qn('w:val'), couleur_clinkast)  # Bleu gras
        r1_size = OxmlElement('w:sz')
        r1_size.set(qn('w:val'), '20')
        r1_props.append(r1_bold)
        r1_props.append(r1_color)
        r1_props.append(r1_size)
        r1_element.append(r1_props)
        
        r1_text = OxmlElement('w:t')
        r1_text.text = categorie
        r1_element.append(r1_text)
        
        p1_element.append(r1_element)
        tc1_element.append(p1_element)
        
        # Cellule 2 - Contenu
        tc2_element = OxmlElement('w:tc')
        tc2_props = OxmlElement('w:tcPr')
        
        tc2_width = OxmlElement('w:tcW')
        tc2_width.set(qn('w:w'), '3000')
        tc2_width.set(qn('w:type'), 'pct')
        tc2_props.append(tc2_width)
        
        # Supprimer complètement les marges internes
        tc2_margins = OxmlElement('w:tcMar')
        tc2_left = OxmlElement('w:left')
        tc2_left.set(qn('w:w'), '0')
        tc2_left.set(qn('w:type'), 'dxa')
        tc2_right = OxmlElement('w:right')
        tc2_right.set(qn('w:w'), '0')
        tc2_right.set(qn('w:type'), 'dxa')
        tc2_top = OxmlElement('w:top')
        tc2_top.set(qn('w:w'), '0')
        tc2_top.set(qn('w:type'), 'dxa')
        tc2_bottom = OxmlElement('w:bottom')
        tc2_bottom.set(qn('w:w'), '120')
        tc2_bottom.set(qn('w:type'), 'dxa')
        tc2_margins.append(tc2_left)
        tc2_margins.append(tc2_right)
        tc2_margins.append(tc2_top)
        tc2_margins.append(tc2_bottom)
        tc2_props.append(tc2_margins)
        
        tc2_element.append(tc2_props)
        
        p2_element = OxmlElement('w:p')
        p2_props = OxmlElement('w:pPr')
        p2_align = OxmlElement('w:jc')
        p2_align.set(qn('w:val'), 'left')
        p2_props.append(p2_align)
        
        # Supprimer complètement toute indentation et espacement
        p2_ind = OxmlElement('w:ind')
        p2_ind.set(qn('w:left'), '0')
        p2_ind.set(qn('w:leftChars'), '0')
        p2_ind.set(qn('w:firstLine'), '0')
        p2_ind.set(qn('w:hanging'), '0')
        p2_props.append(p2_ind)
        
        # Supprimer l'espacement avant et après  
        p2_spacing = OxmlElement('w:spacing')
        p2_spacing.set(qn('w:before'), '0')
        p2_spacing.set(qn('w:after'), '0')
        p2_spacing.set(qn('w:line'), '240')
        p2_spacing.set(qn('w:lineRule'), 'auto')
        p2_props.append(p2_spacing)
        
        p2_element.append(p2_props)
        
        r2_element = OxmlElement('w:r')
        r2_props = OxmlElement('w:rPr')
        r2_size = OxmlElement('w:sz')
        r2_size.set(qn('w:val'), '18')
        r2_color = OxmlElement('w:color')
        r2_color.set(qn('w:val'), couleur_clinkast)  # Bleu normal
        r2_props.append(r2_size)
        r2_props.append(r2_color)
        r2_element.append(r2_props)
        
        r2_text = OxmlElement('w:t')
        r2_text.text = contenu
        r2_element.append(r2_text)
        
        p2_element.append(r2_element)
        tc2_element.append(p2_element)
        
        tr_element.append(tc1_element)
        tr_element.append(tc2_element)

        # Ajouter des propriétés à la ligne
        # tr_props = OxmlElement('w:trPr')
        # tr_height = OxmlElement('w:trHeight')
        # tr_height.set(qn('w:val'), '400')  # Hauteur minimale de ligne
        # tr_height.set(qn('w:hRule'), 'atLeast')
        # tr_props.append(tr_height)
        # tr_element.append(tr_props)
        
        table_element.append(tr_element)
    
    parent_element.insert(paragraph_index + 1, table_element)
    
    return True

def creer_tableau_formation_a_position(doc, paragraph_position, formations_list):
    """
    Crée un tableau de formation à une position spécifique avec style bleu
    """
    if not formations_list or not isinstance(formations_list, list):
        formations_list = [{'annee': '2024', 'intitule': 'Formation à définir'}]
    
    parent_element = paragraph_position._element.getparent()
    paragraph_index = list(parent_element).index(paragraph_position._element)
    
    table_element = parent_element.makeelement(qn('w:tbl'))
    
    tbl_props = OxmlElement('w:tblPr')
    tbl_style = OxmlElement('w:tblStyle')
    tbl_style.set(qn('w:val'), 'TableGrid')
    tbl_props.append(tbl_style)
    
    tbl_width = OxmlElement('w:tblW')
    tbl_width.set(qn('w:w'), '5000')
    tbl_width.set(qn('w:type'), 'pct')
    tbl_props.append(tbl_width)
    
    table_element.append(tbl_props)
    
    couleur_clinkast = "1F4E79"
    
    for formation in formations_list:
        annee = formation.get('annee', 'N/A')
        intitule = formation.get('intitule', 'Formation à définir')
        
        tr_element = OxmlElement('w:tr')
        
        # Cellule 1 - Année
        tc1_element = OxmlElement('w:tc')
        tc1_props = OxmlElement('w:tcPr')
        
        tc1_width = OxmlElement('w:tcW')
        tc1_width.set(qn('w:w'), '1500')
        tc1_width.set(qn('w:type'), 'pct')
        tc1_props.append(tc1_width)
        
        # Supprimer complètement les marges internes
        tc1_margins = OxmlElement('w:tcMar')
        tc1_left = OxmlElement('w:left')
        tc1_left.set(qn('w:w'), '0')
        tc1_left.set(qn('w:type'), 'dxa')
        tc1_right = OxmlElement('w:right')
        tc1_right.set(qn('w:w'), '0')
        tc1_right.set(qn('w:type'), 'dxa')
        tc1_top = OxmlElement('w:top')
        tc1_top.set(qn('w:w'), '0')
        tc1_top.set(qn('w:type'), 'dxa')
        tc1_bottom = OxmlElement('w:bottom')
        tc1_bottom.set(qn('w:w'), '120')
        tc1_bottom.set(qn('w:type'), 'dxa')
        tc1_margins.append(tc1_left)
        tc1_margins.append(tc1_right)
        tc1_margins.append(tc1_top)
        tc1_margins.append(tc1_bottom)
        tc1_props.append(tc1_margins)
        
        tc1_element.append(tc1_props)
        
        p1_element = OxmlElement('w:p')
        p1_props = OxmlElement('w:pPr')
        p1_align = OxmlElement('w:jc')
        p1_align.set(qn('w:val'), 'left')
        p1_props.append(p1_align)
        
        # Supprimer complètement toute indentation et espacement (connaissances - plus strict)
        p1_ind = OxmlElement('w:ind')
        p1_ind.set(qn('w:left'), '0')
        p1_ind.set(qn('w:leftChars'), '0')
        p1_ind.set(qn('w:firstLine'), '0')
        p1_ind.set(qn('w:firstLineChars'), '0')
        p1_ind.set(qn('w:hanging'), '0')
        p1_ind.set(qn('w:hangingChars'), '0')
        p1_ind.set(qn('w:right'), '0')
        p1_ind.set(qn('w:rightChars'), '0')
        p1_props.append(p1_ind)
        
        # Supprimer l'espacement avant et après (plus strict)
        p1_spacing = OxmlElement('w:spacing')
        p1_spacing.set(qn('w:before'), '0')
        p1_spacing.set(qn('w:beforeLines'), '0')
        p1_spacing.set(qn('w:after'), '0')
        p1_spacing.set(qn('w:afterLines'), '0')
        p1_spacing.set(qn('w:line'), '240')
        p1_spacing.set(qn('w:lineRule'), 'auto')
        p1_props.append(p1_spacing)
        
        # Ajouter des marges de paragraphe nulles
        p1_pbdr = OxmlElement('w:pBdr')
        p1_props.append(p1_pbdr)
        
        p1_element.append(p1_props)
        
        r1_element = OxmlElement('w:r')
        r1_props = OxmlElement('w:rPr')
        r1_bold = OxmlElement('w:b')
        r1_color = OxmlElement('w:color')
        r1_color.set(qn('w:val'), couleur_clinkast)  # Bleu gras
        r1_size = OxmlElement('w:sz')
        r1_size.set(qn('w:val'), '20')
        r1_props.append(r1_bold)
        r1_props.append(r1_color)
        r1_props.append(r1_size)
        r1_element.append(r1_props)
        
        r1_text = OxmlElement('w:t')
        r1_text.text = annee
        r1_element.append(r1_text)
        
        p1_element.append(r1_element)
        tc1_element.append(p1_element)
        
        # Cellule 2 - Formation
        tc2_element = OxmlElement('w:tc')
        tc2_props = OxmlElement('w:tcPr')
        
        tc2_width = OxmlElement('w:tcW')
        tc2_width.set(qn('w:w'), '3500')
        tc2_width.set(qn('w:type'), 'pct')
        tc2_props.append(tc2_width)
        
        # Supprimer complètement les marges internes
        tc2_margins = OxmlElement('w:tcMar')
        tc2_left = OxmlElement('w:left')
        tc2_left.set(qn('w:w'), '0')
        tc2_left.set(qn('w:type'), 'dxa')
        tc2_right = OxmlElement('w:right')
        tc2_right.set(qn('w:w'), '0')
        tc2_right.set(qn('w:type'), 'dxa')
        tc2_top = OxmlElement('w:top')
        tc2_top.set(qn('w:w'), '0')
        tc2_top.set(qn('w:type'), 'dxa')
        tc2_bottom = OxmlElement('w:bottom')
        tc2_bottom.set(qn('w:w'), '120')
        tc2_bottom.set(qn('w:type'), 'dxa')
        tc2_margins.append(tc2_left)
        tc2_margins.append(tc2_right)
        tc2_margins.append(tc2_top)
        tc2_margins.append(tc2_bottom)
        tc2_props.append(tc2_margins)
        
        tc2_element.append(tc2_props)
        
        p2_element = OxmlElement('w:p')
        p2_props = OxmlElement('w:pPr')
        p2_align = OxmlElement('w:jc')
        p2_align.set(qn('w:val'), 'left')
        p2_props.append(p2_align)
        
        # Supprimer complètement toute indentation et espacement
        p2_ind = OxmlElement('w:ind')
        p2_ind.set(qn('w:left'), '0')
        p2_ind.set(qn('w:leftChars'), '0')
        p2_ind.set(qn('w:firstLine'), '0')
        p2_ind.set(qn('w:hanging'), '0')
        p2_props.append(p2_ind)
        
        # Supprimer l'espacement avant et après  
        p2_spacing = OxmlElement('w:spacing')
        p2_spacing.set(qn('w:before'), '0')
        p2_spacing.set(qn('w:after'), '0')
        p2_spacing.set(qn('w:line'), '240')
        p2_spacing.set(qn('w:lineRule'), 'auto')
        p2_props.append(p2_spacing)
        
        p2_element.append(p2_props)
        
        r2_element = OxmlElement('w:r')
        r2_props = OxmlElement('w:rPr')
        r2_size = OxmlElement('w:sz')
        r2_size.set(qn('w:val'), '18')
        r2_color = OxmlElement('w:color')
        r2_color.set(qn('w:val'), couleur_clinkast)  # Bleu normal
        r2_props.append(r2_size)
        r2_props.append(r2_color)
        r2_element.append(r2_props)
        
        r2_text = OxmlElement('w:t')
        r2_text.text = intitule
        r2_element.append(r2_text)
        
        p2_element.append(r2_element)
        tc2_element.append(p2_element)
        
        tr_element.append(tc1_element)
        tr_element.append(tc2_element)
        
        table_element.append(tr_element)
    
    parent_element.insert(paragraph_index + 1, table_element)
    
    return True

def creer_tableau_hobbies_a_position(doc, paragraph_position, hobbies_dict):
    """
    Crée un tableau hobbies & divers à une position spécifique avec style bleu
    """
    if not hobbies_dict or not isinstance(hobbies_dict, dict):
        hobbies_dict = {
            'langues': 'Français, Anglais (intermédiaire)',
            'hobbies': 'À définir'
        }
    
    parent_element = paragraph_position._element.getparent()
    paragraph_index = list(parent_element).index(paragraph_position._element)
    
    table_element = parent_element.makeelement(qn('w:tbl'))
    
    tbl_props = OxmlElement('w:tblPr')
    tbl_style = OxmlElement('w:tblStyle')
    tbl_style.set(qn('w:val'), 'TableGrid')
    tbl_props.append(tbl_style)
    
    tbl_width = OxmlElement('w:tblW')
    tbl_width.set(qn('w:w'), '5000')
    tbl_width.set(qn('w:type'), 'pct')
    tbl_props.append(tbl_width)
    
    table_element.append(tbl_props)
    
    couleur_clinkast = "1F4E79"
    
    labels = {
        'langues': 'Langues',
        'hobbies': 'Hobbies'
    }
    
    for key, value in hobbies_dict.items():
        label = labels.get(key, key.capitalize())
        
        tr_element = OxmlElement('w:tr')
        
        # Cellule 1 - Label
        tc1_element = OxmlElement('w:tc')
        tc1_props = OxmlElement('w:tcPr')
        
        tc1_width = OxmlElement('w:tcW')
        tc1_width.set(qn('w:w'), '1500')
        tc1_width.set(qn('w:type'), 'pct')
        tc1_props.append(tc1_width)
        
        # Supprimer complètement les marges internes (tableau hobbies - cellule 1)
        tc1_margins = OxmlElement('w:tcMar')
        tc1_left = OxmlElement('w:left')
        tc1_left.set(qn('w:w'), '0')
        tc1_left.set(qn('w:type'), 'dxa')
        tc1_right = OxmlElement('w:right')
        tc1_right.set(qn('w:w'), '0')
        tc1_right.set(qn('w:type'), 'dxa')
        tc1_top = OxmlElement('w:top')
        tc1_top.set(qn('w:w'), '0')
        tc1_top.set(qn('w:type'), 'dxa')
        tc1_bottom = OxmlElement('w:bottom')
        tc1_bottom.set(qn('w:w'), '120')
        tc1_bottom.set(qn('w:type'), 'dxa')
        tc1_margins.append(tc1_left)
        tc1_margins.append(tc1_right)
        tc1_margins.append(tc1_top)
        tc1_margins.append(tc1_bottom)
        tc1_props.append(tc1_margins)
        
        tc1_element.append(tc1_props)
        
        p1_element = OxmlElement('w:p')
        p1_props = OxmlElement('w:pPr')
        p1_align = OxmlElement('w:jc')
        p1_align.set(qn('w:val'), 'left')
        p1_props.append(p1_align)
        
        # Supprimer complètement toute indentation et espacement (hobbies)
        p1_ind = OxmlElement('w:ind')
        p1_ind.set(qn('w:left'), '0')
        p1_ind.set(qn('w:leftChars'), '0')
        p1_ind.set(qn('w:firstLine'), '0')
        p1_ind.set(qn('w:hanging'), '0')
        p1_props.append(p1_ind)
        
        # Supprimer l'espacement avant et après (hobbies)
        p1_spacing = OxmlElement('w:spacing')
        p1_spacing.set(qn('w:before'), '0')
        p1_spacing.set(qn('w:after'), '0')
        p1_spacing.set(qn('w:line'), '240')
        p1_spacing.set(qn('w:lineRule'), 'auto')
        p1_props.append(p1_spacing)
        
        p1_element.append(p1_props)
        
        r1_element = OxmlElement('w:r')
        r1_props = OxmlElement('w:rPr')
        r1_bold = OxmlElement('w:b')
        r1_color = OxmlElement('w:color')
        r1_color.set(qn('w:val'), couleur_clinkast)  # Bleu gras
        r1_size = OxmlElement('w:sz')
        r1_size.set(qn('w:val'), '20')
        r1_props.append(r1_bold)
        r1_props.append(r1_color)
        r1_props.append(r1_size)
        r1_element.append(r1_props)
        
        r1_text = OxmlElement('w:t')
        r1_text.text = f"{label} :"
        r1_element.append(r1_text)
        
        p1_element.append(r1_element)
        tc1_element.append(p1_element)
        
        # Cellule 2 - Contenu
        tc2_element = OxmlElement('w:tc')
        tc2_props = OxmlElement('w:tcPr')
        
        tc2_width = OxmlElement('w:tcW')
        tc2_width.set(qn('w:w'), '3500')
        tc2_width.set(qn('w:type'), 'pct')
        tc2_props.append(tc2_width)
        
        # Supprimer complètement les marges internes (tableau hobbies - cellule 2)
        tc2_margins = OxmlElement('w:tcMar')
        tc2_left = OxmlElement('w:left')
        tc2_left.set(qn('w:w'), '0')
        tc2_left.set(qn('w:type'), 'dxa')
        tc2_right = OxmlElement('w:right')
        tc2_right.set(qn('w:w'), '0')
        tc2_right.set(qn('w:type'), 'dxa')
        tc2_top = OxmlElement('w:top')
        tc2_top.set(qn('w:w'), '0')
        tc2_top.set(qn('w:type'), 'dxa')
        tc2_bottom = OxmlElement('w:bottom')
        tc2_bottom.set(qn('w:w'), '120')
        tc2_bottom.set(qn('w:type'), 'dxa')
        tc2_margins.append(tc2_left)
        tc2_margins.append(tc2_right)
        tc2_margins.append(tc2_top)
        tc2_margins.append(tc2_bottom)
        tc2_props.append(tc2_margins)
        
        tc2_element.append(tc2_props)
        
        p2_element = OxmlElement('w:p')
        p2_props = OxmlElement('w:pPr')
        p2_align = OxmlElement('w:jc')
        p2_align.set(qn('w:val'), 'left')
        p2_props.append(p2_align)
        
        # Supprimer complètement toute indentation et espacement (hobbies cell 2)
        p2_ind = OxmlElement('w:ind')
        p2_ind.set(qn('w:left'), '0')
        p2_ind.set(qn('w:leftChars'), '0')
        p2_ind.set(qn('w:firstLine'), '0')
        p2_ind.set(qn('w:hanging'), '0')
        p2_props.append(p2_ind)
        
        # Supprimer l'espacement avant et après (hobbies cell 2)
        p2_spacing = OxmlElement('w:spacing')
        p2_spacing.set(qn('w:before'), '0')
        p2_spacing.set(qn('w:after'), '0')
        p2_spacing.set(qn('w:line'), '240')
        p2_spacing.set(qn('w:lineRule'), 'auto')
        p2_props.append(p2_spacing)
        
        p2_element.append(p2_props)
        
        r2_element = OxmlElement('w:r')
        r2_props = OxmlElement('w:rPr')
        r2_size = OxmlElement('w:sz')
        r2_size.set(qn('w:val'), '18')
        r2_color = OxmlElement('w:color')
        r2_color.set(qn('w:val'), couleur_clinkast)  # Bleu normal
        r2_props.append(r2_size)
        r2_props.append(r2_color)
        r2_element.append(r2_props)
        
        r2_text = OxmlElement('w:t')
        r2_text.text = value
        r2_element.append(r2_text)
        
        p2_element.append(r2_element)
        tc2_element.append(p2_element)
        
        tr_element.append(tc1_element)
        tr_element.append(tc2_element)
        
        table_element.append(tr_element)
    
    parent_element.insert(paragraph_index + 1, table_element)
    
    return True

def creer_section_experiences_a_position(doc, paragraph_position, experiences_list):
    """
    Crée une section expériences formatée à une position spécifique
    """
    if not experiences_list or not isinstance(experiences_list, list):
        return False
    
    parent_element = paragraph_position._element.getparent()
    paragraph_index = list(parent_element).index(paragraph_position._element)
    
    current_index = paragraph_index + 1
    couleur_clinkast = "1F4E79"
    
    for i, exp in enumerate(experiences_list):
        # Créer un tableau pour l'en-tête (comme dans l'image)
        table_entete = parent_element.makeelement(qn('w:tbl'))
        
        # Propriétés du tableau en-tête
        tbl_entete_props = OxmlElement('w:tblPr')
        tbl_entete_style = OxmlElement('w:tblStyle')
        tbl_entete_style.set(qn('w:val'), 'TableGrid')
        tbl_entete_props.append(tbl_entete_style)
        
        tbl_entete_width = OxmlElement('w:tblW')
        tbl_entete_width.set(qn('w:w'), '5000')
        tbl_entete_width.set(qn('w:type'), 'pct')
        tbl_entete_props.append(tbl_entete_width)
        
        # Espacement avant le tableau
        tbl_entete_spacing = OxmlElement('w:spacing')
        tbl_entete_spacing.set(qn('w:before'), '240')
        tbl_entete_props.append(tbl_entete_spacing)
        
        table_entete.append(tbl_entete_props)
        
        # Créer la ligne d'en-tête
        tr_entete = OxmlElement('w:tr')
        
        # Une seule cellule qui couvre toute la largeur
        tc_entete = OxmlElement('w:tc')
        tc_entete_props = OxmlElement('w:tcPr')
        
        # Largeur de la cellule
        tc_entete_width = OxmlElement('w:tcW')
        tc_entete_width.set(qn('w:w'), '5000')
        tc_entete_width.set(qn('w:type'), 'pct')
        tc_entete_props.append(tc_entete_width)
        
        # Couleur de fond marron/doré (comme dans l'image)
        shd_entete = OxmlElement('w:shd')
        shd_entete.set(qn('w:val'), 'clear')
        shd_entete.set(qn('w:color'), 'auto')
        shd_entete.set(qn('w:fill'), 'B8860B')  # DarkGoldenRod
        tc_entete_props.append(shd_entete)
        
        # Alignement vertical
        vAlign_entete = OxmlElement('w:vAlign')
        vAlign_entete.set(qn('w:val'), 'center')
        tc_entete_props.append(vAlign_entete)
        
        tc_entete.append(tc_entete_props)
        
        # Paragraphe dans la cellule
        p_entete = OxmlElement('w:p')
        p_entete_props = OxmlElement('w:pPr')
        p_entete_align = OxmlElement('w:jc')
        p_entete_align.set(qn('w:val'), 'left')  # Centré à gauche
        p_entete_props.append(p_entete_align)
        
        # Supprimer les marges internes
        p_entete_ind = OxmlElement('w:ind')
        p_entete_ind.set(qn('w:left'), '0')
        p_entete_ind.set(qn('w:right'), '0')
        p_entete_props.append(p_entete_ind)
        
        p_entete.append(p_entete_props)
        
        # Texte période
        r_periode = OxmlElement('w:r')
        r_periode_props = OxmlElement('w:rPr')
        r_periode_bold = OxmlElement('w:b')
        r_periode_color = OxmlElement('w:color')
        r_periode_color.set(qn('w:val'), 'FFFFFF')  # Blanc
        r_periode_size = OxmlElement('w:sz')
        r_periode_size.set(qn('w:val'), '20')  # Réduit à 10pt
        r_periode_props.append(r_periode_bold)
        r_periode_props.append(r_periode_color)
        r_periode_props.append(r_periode_size)
        r_periode.append(r_periode_props)
        
        r_periode_text = OxmlElement('w:t')
        r_periode_text.text = exp.get('periode', 'Période')
        r_periode.append(r_periode_text)
        
        # Tiret et espace après la période
        r_tiret = OxmlElement('w:r')
        r_tiret_props = OxmlElement('w:rPr')
        r_tiret_bold = OxmlElement('w:b')
        r_tiret_color = OxmlElement('w:color')
        r_tiret_color.set(qn('w:val'), 'FFFFFF')
        r_tiret_size = OxmlElement('w:sz')
        r_tiret_size.set(qn('w:val'), '20')
        r_tiret_props.append(r_tiret_bold)
        r_tiret_props.append(r_tiret_color)
        r_tiret_props.append(r_tiret_size)
        r_tiret.append(r_tiret_props)
        
        r_tiret_text = OxmlElement('w:t')
        r_tiret_text.text = ' - '
        r_tiret.append(r_tiret_text)
        
        # Texte titre
        r_titre = OxmlElement('w:r')
        r_titre_props = OxmlElement('w:rPr')
        r_titre_bold = OxmlElement('w:b')
        r_titre_color = OxmlElement('w:color')
        r_titre_color.set(qn('w:val'), 'FFFFFF')
        r_titre_size = OxmlElement('w:sz')
        r_titre_size.set(qn('w:val'), '20')  # Réduit à 10pt
        r_titre_props.append(r_titre_bold)
        r_titre_props.append(r_titre_color)
        r_titre_props.append(r_titre_size)
        r_titre.append(r_titre_props)
        
        r_titre_text = OxmlElement('w:t')
        r_titre_text.text = exp.get('titre', 'Titre')
        r_titre.append(r_titre_text)
        
        # Séparateur 2
        r_sep2 = OxmlElement('w:r')
        r_sep2_props = OxmlElement('w:rPr')
        r_sep2_bold = OxmlElement('w:b')
        r_sep2_color = OxmlElement('w:color')
        r_sep2_color.set(qn('w:val'), 'FFFFFF')
        r_sep2_size = OxmlElement('w:sz')
        r_sep2_size.set(qn('w:val'), '20')
        r_sep2_props.append(r_sep2_bold)
        r_sep2_props.append(r_sep2_color)
        r_sep2_props.append(r_sep2_size)
        r_sep2.append(r_sep2_props)
        
        r_sep2_text = OxmlElement('w:t')
        r_sep2_text.text = ' - '
        r_sep2.append(r_sep2_text)
        
        # Texte entreprise
        r_entreprise = OxmlElement('w:r')
        r_entreprise_props = OxmlElement('w:rPr')
        r_entreprise_bold = OxmlElement('w:b')
        r_entreprise_color = OxmlElement('w:color')
        r_entreprise_color.set(qn('w:val'), 'FFFFFF')
        r_entreprise_size = OxmlElement('w:sz')
        r_entreprise_size.set(qn('w:val'), '20')  # Réduit à 10pt
        r_entreprise_props.append(r_entreprise_bold)
        r_entreprise_props.append(r_entreprise_color)
        r_entreprise_props.append(r_entreprise_size)
        r_entreprise.append(r_entreprise_props)
        
        r_entreprise_text = OxmlElement('w:t')
        r_entreprise_text.text = exp.get('entreprise', 'Entreprise')
        r_entreprise.append(r_entreprise_text)
        
        # Ajouter tous les runs au paragraphe
        p_entete.append(r_periode)
        p_entete.append(r_tiret)
        p_entete.append(r_titre)
        p_entete.append(r_sep2)
        p_entete.append(r_entreprise)
        
        # Ajouter le paragraphe à la cellule
        tc_entete.append(p_entete)
        
        # Ajouter la cellule à la ligne
        tr_entete.append(tc_entete)
        
        # Ajouter la ligne au tableau
        table_entete.append(tr_entete)
        
        # Insérer le tableau dans le document
        parent_element.insert(current_index, table_entete)
        current_index += 1
        
        # Ligne vide
        p_vide = parent_element.makeelement(qn('w:p'))
        parent_element.insert(current_index, p_vide)
        current_index += 1
        
        # Responsabilité (BLEU ET GRAS SANS INDENTATION)
        p_resp = parent_element.makeelement(qn('w:p'))
        
        # Propriétés sans indentation
        p_resp_props = OxmlElement('w:pPr')
        p_resp_align = OxmlElement('w:jc')
        p_resp_align.set(qn('w:val'), 'left')
        p_resp_props.append(p_resp_align)
        
        p_resp_ind = OxmlElement('w:ind')
        p_resp_ind.set(qn('w:left'), '0')
        p_resp_ind.set(qn('w:leftChars'), '0')
        p_resp_ind.set(qn('w:firstLine'), '0')
        p_resp_ind.set(qn('w:firstLineChars'), '0')
        p_resp_ind.set(qn('w:hanging'), '0')
        p_resp_ind.set(qn('w:hangingChars'), '0')
        p_resp_ind.set(qn('w:right'), '0')
        p_resp_ind.set(qn('w:rightChars'), '0')
        p_resp_props.append(p_resp_ind)
        
        p_resp_spacing = OxmlElement('w:spacing')
        p_resp_spacing.set(qn('w:before'), '0')
        p_resp_spacing.set(qn('w:beforeLines'), '0')
        p_resp_spacing.set(qn('w:after'), '0')
        p_resp_spacing.set(qn('w:afterLines'), '0')
        p_resp_spacing.set(qn('w:line'), '240')
        p_resp_spacing.set(qn('w:lineRule'), 'auto')
        p_resp_props.append(p_resp_spacing)
        
        p_resp_pbdr = OxmlElement('w:pBdr')
        p_resp_props.append(p_resp_pbdr)
        
        p_resp.append(p_resp_props)
        
        r_resp_label = OxmlElement('w:r')
        r_resp_label_props = OxmlElement('w:rPr')
        r_resp_label_bold = OxmlElement('w:b')
        r_resp_label_color = OxmlElement('w:color')
        r_resp_label_color.set(qn('w:val'), couleur_clinkast)  # BLEU
        r_resp_label_size = OxmlElement('w:sz')
        r_resp_label_size.set(qn('w:val'), '20')
        r_resp_label_props.append(r_resp_label_bold)
        r_resp_label_props.append(r_resp_label_color)
        r_resp_label_props.append(r_resp_label_size)
        r_resp_label.append(r_resp_label_props)
        
        r_resp_label_text = OxmlElement('w:t')
        r_resp_label_text.text = 'Responsabilité : '
        r_resp_label.append(r_resp_label_text)
        
        r_resp_content = OxmlElement('w:r')
        r_resp_content_props = OxmlElement('w:rPr')
        r_resp_content_size = OxmlElement('w:sz')
        r_resp_content_size.set(qn('w:val'), '20')
        r_resp_content_props.append(r_resp_content_size)
        r_resp_content.append(r_resp_content_props)
        
        r_resp_content_text = OxmlElement('w:t')
        r_resp_content_text.text = exp.get('responsabilites', 'Responsabilités à définir')
        r_resp_content.append(r_resp_content_text)
        
        p_resp.append(r_resp_label)
        p_resp.append(r_resp_content)
        
        parent_element.insert(current_index, p_resp)
        current_index += 1
        
        # Ligne vide
        p_vide2 = parent_element.makeelement(qn('w:p'))
        parent_element.insert(current_index, p_vide2)
        current_index += 1
        
        # Titre Réalisations (BLEU ET GRAS SANS INDENTATION)
        p_real_titre = parent_element.makeelement(qn('w:p'))
        
        # Propriétés sans indentation
        p_real_titre_props = OxmlElement('w:pPr')
        p_real_titre_align = OxmlElement('w:jc')
        p_real_titre_align.set(qn('w:val'), 'left')
        p_real_titre_props.append(p_real_titre_align)
        
        p_real_titre_ind = OxmlElement('w:ind')
        p_real_titre_ind.set(qn('w:left'), '0')
        p_real_titre_ind.set(qn('w:leftChars'), '0')
        p_real_titre_ind.set(qn('w:firstLine'), '0')
        p_real_titre_ind.set(qn('w:firstLineChars'), '0')
        p_real_titre_ind.set(qn('w:hanging'), '0')
        p_real_titre_ind.set(qn('w:hangingChars'), '0')
        p_real_titre_ind.set(qn('w:right'), '0')
        p_real_titre_ind.set(qn('w:rightChars'), '0')
        p_real_titre_props.append(p_real_titre_ind)
        
        p_real_titre_spacing = OxmlElement('w:spacing')
        p_real_titre_spacing.set(qn('w:before'), '0')
        p_real_titre_spacing.set(qn('w:beforeLines'), '0')
        p_real_titre_spacing.set(qn('w:after'), '0')
        p_real_titre_spacing.set(qn('w:afterLines'), '0')
        p_real_titre_spacing.set(qn('w:line'), '240')
        p_real_titre_spacing.set(qn('w:lineRule'), 'auto')
        p_real_titre_props.append(p_real_titre_spacing)
        
        p_real_titre_pbdr = OxmlElement('w:pBdr')
        p_real_titre_props.append(p_real_titre_pbdr)
        
        p_real_titre.append(p_real_titre_props)
        
        r_real_titre = OxmlElement('w:r')
        r_real_titre_props = OxmlElement('w:rPr')
        r_real_titre_bold = OxmlElement('w:b')
        r_real_titre_color = OxmlElement('w:color')
        r_real_titre_color.set(qn('w:val'), couleur_clinkast)  # BLEU
        r_real_titre_size = OxmlElement('w:sz')
        r_real_titre_size.set(qn('w:val'), '20')
        r_real_titre_props.append(r_real_titre_bold)
        r_real_titre_props.append(r_real_titre_color)
        r_real_titre_props.append(r_real_titre_size)
        r_real_titre.append(r_real_titre_props)
        
        r_real_titre_text = OxmlElement('w:t')
        r_real_titre_text.text = 'Réalisations'
        r_real_titre.append(r_real_titre_text)
        p_real_titre.append(r_real_titre)
        
        parent_element.insert(current_index, p_real_titre)
        current_index += 1
        
        # Liste des réalisations (indentées et en noir)
        realisations = exp.get('realisations', ['Réalisation à définir'])
        if isinstance(realisations, list):
            for realisation in realisations:
                p_real_item = parent_element.makeelement(qn('w:p'))
                p_real_item_props = OxmlElement('w:pPr')
                p_real_item_ind = OxmlElement('w:ind')
                p_real_item_ind.set(qn('w:left'), '360')  # Indentation normale
                p_real_item_props.append(p_real_item_ind)
                p_real_item.append(p_real_item_props)
                
                r_real_item = OxmlElement('w:r')
                r_real_item_props = OxmlElement('w:rPr')
                r_real_item_size = OxmlElement('w:sz')
                r_real_item_size.set(qn('w:val'), '18')
                r_real_item_props.append(r_real_item_size)
                # Pas de couleur spécifiée = noir par défaut
                r_real_item.append(r_real_item_props)
                
                r_real_item_text = OxmlElement('w:t')
                r_real_item_text.text = f"• {realisation}"
                r_real_item.append(r_real_item_text)
                p_real_item.append(r_real_item)
                
                parent_element.insert(current_index, p_real_item)
                current_index += 1
        
        # Ligne vide
        p_vide3 = parent_element.makeelement(qn('w:p'))
        parent_element.insert(current_index, p_vide3)
        current_index += 1
        
        # Environnement (BLEU ET GRAS SANS INDENTATION)
        p_env = parent_element.makeelement(qn('w:p'))
        
        # Propriétés sans indentation
        p_env_props = OxmlElement('w:pPr')
        p_env_align = OxmlElement('w:jc')
        p_env_align.set(qn('w:val'), 'left')
        p_env_props.append(p_env_align)
        
        p_env_ind = OxmlElement('w:ind')
        p_env_ind.set(qn('w:left'), '0')
        p_env_ind.set(qn('w:leftChars'), '0')
        p_env_ind.set(qn('w:firstLine'), '0')
        p_env_ind.set(qn('w:firstLineChars'), '0')
        p_env_ind.set(qn('w:hanging'), '0')
        p_env_ind.set(qn('w:hangingChars'), '0')
        p_env_ind.set(qn('w:right'), '0')
        p_env_ind.set(qn('w:rightChars'), '0')
        p_env_props.append(p_env_ind)
        
        p_env_spacing = OxmlElement('w:spacing')
        p_env_spacing.set(qn('w:before'), '0')
        p_env_spacing.set(qn('w:beforeLines'), '0')
        p_env_spacing.set(qn('w:after'), '0')
        p_env_spacing.set(qn('w:afterLines'), '0')
        p_env_spacing.set(qn('w:line'), '240')
        p_env_spacing.set(qn('w:lineRule'), 'auto')
        p_env_props.append(p_env_spacing)
        
        p_env_pbdr = OxmlElement('w:pBdr')
        p_env_props.append(p_env_pbdr)
        
        p_env.append(p_env_props)
        
        r_env_label = OxmlElement('w:r')
        r_env_label_props = OxmlElement('w:rPr')
        r_env_label_bold = OxmlElement('w:b')
        r_env_label_color = OxmlElement('w:color')
        r_env_label_color.set(qn('w:val'), couleur_clinkast)  # BLEU
        r_env_label_underline = OxmlElement('w:u')  # Souligner Environnement
        r_env_label_underline.set(qn('w:val'), 'single')
        r_env_label_size = OxmlElement('w:sz')
        r_env_label_size.set(qn('w:val'), '20')
        r_env_label_props.append(r_env_label_bold)
        r_env_label_props.append(r_env_label_color)
        r_env_label_props.append(r_env_label_underline)
        r_env_label_props.append(r_env_label_size)
        r_env_label.append(r_env_label_props)
        
        r_env_label_text = OxmlElement('w:t')
        r_env_label_text.text = 'Environnement : '
        r_env_label.append(r_env_label_text)
        
        r_env_content = OxmlElement('w:r')
        r_env_content_props = OxmlElement('w:rPr')
        r_env_content_size = OxmlElement('w:sz')
        r_env_content_size.set(qn('w:val'), '18')
        r_env_content_props.append(r_env_content_size)
        r_env_content.append(r_env_content_props)
        
        r_env_content_text = OxmlElement('w:t')
        r_env_content_text.text = exp.get('environnement', 'Environnement à définir')
        r_env_content.append(r_env_content_text)
        
        p_env.append(r_env_label)
        p_env.append(r_env_content)
        
        parent_element.insert(current_index, p_env)
        current_index += 1
        
        # Espace entre expériences
        if i < len(experiences_list) - 1:
            p_separation = parent_element.makeelement(qn('w:p'))
            p_sep_props = OxmlElement('w:pPr')
            p_sep_spacing = OxmlElement('w:spacing')
            p_sep_spacing.set(qn('w:after'), '240')
            p_sep_props.append(p_sep_spacing)
            p_separation.append(p_sep_props)
            parent_element.insert(current_index, p_separation)
            current_index += 1
    
    return True

def remplacer_texte_paragraph_avec_style(paragraph, ancien_texte, nouveau_texte, appliquer_style_bleu=False):
    """
    Remplace du texte dans un paragraphe en préservant le formatage
    Avec option pour appliquer un style bleu et gras
    """
    if ancien_texte in paragraph.text:
        for run in paragraph.runs:
            if ancien_texte in run.text:
                run.text = run.text.replace(ancien_texte, nouveau_texte)
                if appliquer_style_bleu:
                    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)  # Bleu Clinkast
                    run.bold = True
                    run.font.size = Pt(12)  # Taille légèrement plus grande
                return True
        
        texte_complet = paragraph.text
        if ancien_texte in texte_complet:
            paragraph.clear()
            nouveau_run = paragraph.add_run(texte_complet.replace(ancien_texte, nouveau_texte))
            if appliquer_style_bleu:
                nouveau_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)  # Bleu Clinkast
                nouveau_run.bold = True
                nouveau_run.font.size = Pt(12)  # Taille légèrement plus grande
            return True
    
    return False

def remplacer_placeholders(doc, data):
    """
    Remplace tous les placeholders dans le document avec les données
    """
    # Remplacements simples avec styles spéciaux
    remplacements_avec_style = {
        '{{nom_consultant}}': (data.get('nom_consultant', 'Nom du consultant'), True),
        '{{titre_du_poste}}': (data.get('titre_du_poste', 'Titre du poste'), True)
    }
    
    # Remplacements simples normaux
    remplacements_normaux = {
        '{{mois_debut_experience}}': data.get('mois_debut_experience', 'Date'),
        '{{nom_entreprise}}': data.get('nom_entreprise', 'Entreprise')
    }
    
    # Variables de contrôle
    tableau_connaissances_cree = False
    tableau_formation_cree = False  
    tableau_hobbies_cree = False
    tableau_experiences_cree = False
    
    # Première passe - traiter les tableaux spéciaux
    paragraphs_to_process = list(doc.paragraphs)
    
    for paragraph in paragraphs_to_process:
        if '{{tableau_connaissances}}' in paragraph.text and not tableau_connaissances_cree:
            connaissances_dict = data.get('connaissances', {})
            paragraph.clear()
            creer_tableau_connaissances_a_position(doc, paragraph, connaissances_dict)
            tableau_connaissances_cree = True
            continue
            
        elif '{{tableau_formation}}' in paragraph.text and not tableau_formation_cree:
            formations_list = data.get('formations', [])
            paragraph.clear()
            creer_tableau_formation_a_position(doc, paragraph, formations_list)
            tableau_formation_cree = True
            continue
            
        elif '{{tableau_hobbies}}' in paragraph.text and not tableau_hobbies_cree:
            hobbies_dict = data.get('hobbies_divers', {})
            paragraph.clear()
            creer_tableau_hobbies_a_position(doc, paragraph, hobbies_dict)
            tableau_hobbies_cree = True
            continue
            
        elif '{{tableau_experiences}}' in paragraph.text and not tableau_experiences_cree:
            experiences_list = data.get('experiences', [])
            paragraph.clear()
            creer_section_experiences_a_position(doc, paragraph, experiences_list)
            tableau_experiences_cree = True
            continue
    
    # Deuxième passe - autres remplacements
    for paragraph in doc.paragraphs:
        # Sécurité pour tableaux non traités
        if '{{tableau_formation}}' in paragraph.text and not tableau_formation_cree:
            formations_list = data.get('formations', [])
            remplacer_texte_paragraph(paragraph, '{{tableau_formation}}', '')
            creer_tableau_formation_a_position(doc, paragraph, formations_list)
            tableau_formation_cree = True
            continue
        
        if '{{tableau_hobbies}}' in paragraph.text and not tableau_hobbies_cree:
            hobbies_dict = data.get('hobbies_divers', {})
            remplacer_texte_paragraph(paragraph, '{{tableau_hobbies}}', '')
            creer_tableau_hobbies_a_position(doc, paragraph, hobbies_dict)
            tableau_hobbies_cree = True
            continue
        
        if '{{tableau_connaissances}}' in paragraph.text and not tableau_connaissances_cree:
            connaissances_dict = data.get('connaissances', {})
            remplacer_texte_paragraph(paragraph, '{{tableau_connaissances}}', '')
            creer_tableau_connaissances_a_position(doc, paragraph, connaissances_dict)
            tableau_connaissances_cree = True
            continue
        
        if '{{tableau_experiences}}' in paragraph.text and not tableau_experiences_cree:
            experiences_list = data.get('experiences', [])
            remplacer_texte_paragraph(paragraph, '{{tableau_experiences}}', '')
            creer_section_experiences_a_position(doc, paragraph, experiences_list)
            tableau_experiences_cree = True
            continue
        
        # Points forts en liste
        if '{{points_forts}}' in paragraph.text:
            points_forts = data.get('points_forts', [])
            if isinstance(points_forts, list) and points_forts:
                # Appliquer le style bleu et gras pour le premier item
                remplacer_texte_paragraph(paragraph, '{{points_forts}}', f"• {points_forts[0]}")
                # Configurer le style du paragraphe modifié
                for run in paragraph.runs:
                    if run.text and '•' in run.text:
                        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)  # Bleu Clinkast
                        run.bold = True
                        run.font.size = Pt(10)
                
                if len(points_forts) > 1:
                    ajouter_liste_paragraphes(doc, paragraph, points_forts[1:], style_bullet=True)
            else:
                remplacer_texte_paragraph(paragraph, '{{points_forts}}', 'Points forts à définir')
            continue
        
        # Niveaux d'intervention en liste
        if '{{niveaux_intervention}}' in paragraph.text:
            niveaux = data.get('niveaux_intervention', [])
            if isinstance(niveaux, list) and niveaux:
                remplacer_texte_paragraph(paragraph, '{{niveaux_intervention}}', f"• {niveaux[0]}")
                # Configurer le style du paragraphe modifié
                for run in paragraph.runs:
                    if run.text and '•' in run.text:
                        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)  # Bleu Clinkast
                        run.bold = True
                        run.font.size = Pt(10)
                
                if len(niveaux) > 1:
                    ajouter_liste_paragraphes(doc, paragraph, niveaux[1:], style_bullet=True)
            else:
                remplacer_texte_paragraph(paragraph, '{{niveaux_intervention}}', 'Niveaux d\'intervention à définir')
            continue
        
        # Remplacements avec style bleu (NOUVEAU)
        for placeholder, (valeur, appliquer_style) in remplacements_avec_style.items():
            if placeholder in paragraph.text:
                remplacer_texte_paragraph_avec_style(paragraph, placeholder, valeur, appliquer_style)
        
        # Remplacements simples normaux
        for placeholder, valeur in remplacements_normaux.items():
            if placeholder in paragraph.text:
                remplacer_texte_paragraph(paragraph, placeholder, valeur)
    
    # Traiter aussi les tableaux existants
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # Remplacements avec style bleu dans les tableaux
                    for placeholder, (valeur, appliquer_style) in remplacements_avec_style.items():
                        if placeholder in paragraph.text:
                            remplacer_texte_paragraph_avec_style(paragraph, placeholder, valeur, appliquer_style)
                    
                    # Remplacements normaux dans les tableaux
                    for placeholder, valeur in remplacements_normaux.items():
                        if placeholder in paragraph.text:
                            remplacer_texte_paragraph(paragraph, placeholder, valeur)
                            

def ajouter_entete_a_partir_page_2_preserve_template(doc, nom_consultant):
    """
    Ajoute une en-tête à droite à partir de la deuxième page
    SANS supprimer le contenu existant du template
    Version avec mise en forme améliorée : bleu, gras, plus grande
    POSITIONNÉE TOUT EN HAUT
    """
    from docx.oxml.shared import OxmlElement, qn
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    for section in doc.sections:
        # Activer l'en-tête différente pour la première page
        section.different_first_page_header_footer = True
        
        # Pour l'en-tête des pages suivantes (2+)
        header = section.header
        
        # CORRECTION : Insérer le nouveau paragraphe EN PREMIER (index 0)
        # au lieu de l'ajouter à la fin
        parent_element = header._element
        
        # Créer le nouvel élément paragraphe
        new_para_element = parent_element.makeelement(qn('w:p'))
        
        # Propriétés du paragraphe - aligné à droite, sans espacement
        p_props = OxmlElement('w:pPr')
        
        # Alignement à droite
        p_align = OxmlElement('w:jc')
        p_align.set(qn('w:val'), 'right')
        p_props.append(p_align)
        
        # Réduire l'espacement au maximum
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:beforeLines'), '0')
        spacing.set(qn('w:after'), '120')  # Petit espacement après pour séparer du contenu existant
        spacing.set(qn('w:afterLines'), '0')
        spacing.set(qn('w:line'), '240')
        spacing.set(qn('w:lineRule'), 'auto')
        p_props.append(spacing)
        
        new_para_element.append(p_props)
        
        # Créer le premier run avec le texte principal
        run1_element = OxmlElement('w:r')
        run1_props = OxmlElement('w:rPr')
        
        # Propriétés de formatage : Calibri, 11pt, bleu, gras
        run1_font = OxmlElement('w:rFonts')
        run1_font.set(qn('w:ascii'), 'Calibri')
        run1_font.set(qn('w:hAnsi'), 'Calibri')
        run1_props.append(run1_font)
        
        run1_size = OxmlElement('w:sz')
        run1_size.set(qn('w:val'), '22')  # 11pt = 22 en half-points
        run1_props.append(run1_size)
        
        run1_color = OxmlElement('w:color')
        run1_color.set(qn('w:val'), '1F4E79')  # Bleu Clinkast
        run1_props.append(run1_color)
        
        run1_bold = OxmlElement('w:b')
        run1_props.append(run1_bold)
        
        run1_element.append(run1_props)
        
        # Texte du premier run
        run1_text = OxmlElement('w:t')
        run1_text.text = f"Dossier de compétences {nom_consultant} / page "
        run1_element.append(run1_text)
        
        new_para_element.append(run1_element)
        
        # Créer le deuxième run avec le champ PAGE
        run2_element = OxmlElement('w:r')
        run2_props = OxmlElement('w:rPr')
        
        # Mêmes propriétés de formatage
        run2_font = OxmlElement('w:rFonts')
        run2_font.set(qn('w:ascii'), 'Calibri')
        run2_font.set(qn('w:hAnsi'), 'Calibri')
        run2_props.append(run2_font)
        
        run2_size = OxmlElement('w:sz')
        run2_size.set(qn('w:val'), '22')  # 11pt
        run2_props.append(run2_size)
        
        run2_color = OxmlElement('w:color')
        run2_color.set(qn('w:val'), '1F4E79')  # Bleu Clinkast
        run2_props.append(run2_color)
        
        run2_bold = OxmlElement('w:b')
        run2_props.append(run2_bold)
        
        run2_element.append(run2_props)
        
        # Champs PAGE
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run2_element.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        run2_element.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run2_element.append(fldChar2)
        
        new_para_element.append(run2_element)
        
        # INSÉRER en première position (index 0) pour être tout en haut
        parent_element.insert(0, new_para_element)


def generer_cv_depuis_template_avec_entete_preserve(template_file, data):
    """
    Version qui préserve tout le contenu du template
    """
    try:
        # Charger le document SANS modification
        doc = Document(template_file)
        
        # D'abord remplir les placeholders normalement
        remplacer_placeholders(doc, data)
        
        # ENSUITE ajouter l'en-tête personnalisée
        nom_consultant = data.get('nom_consultant', 'Nom du consultant')
        ajouter_entete_a_partir_page_2_preserve_template(doc, nom_consultant)
        
        return doc
        
    except Exception as e:
        st.error(f"Erreur lors du traitement du template avec en-tête: {str(e)}")
        import traceback
        st.error(f"Détails: {traceback.format_exc()}")
        return None

# def generer_cv_depuis_template(template_file, data):
#     """
#     Génère un CV en remplissant un template Word
#     """
#     try:
#         doc = Document(template_file)
#         remplacer_placeholders(doc, data)
#         return doc
#     except Exception as e:
#         st.error(f"Erreur lors du traitement du template: {str(e)}")
#         return None