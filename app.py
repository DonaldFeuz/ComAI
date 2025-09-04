import streamlit as st
import json
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import io
import base64

# Configuration de la page
st.set_page_config(
    page_title="Générateur de CV Clinkast - Template",
    page_icon="📄",
    layout="wide"
)

def extraire_contenu_json(reponse_ia):
    """
    Extrait le contenu JSON de la réponse de l'IA
    """
    # Recherche du pattern JSON dans la chaîne de caractères
    pattern = r'content=\'({.*?})\''
    match = re.search(pattern, str(reponse_ia), re.DOTALL)
    
    if match:
        json_str = match.group(1)
        
        # Nettoyage approfondi des échappements
        json_str = json_str.replace('\\\\', '\\')
        json_str = json_str.replace("\\'", "'")
        json_str = json_str.replace('\\n', '\n')
        json_str = re.sub(r'\\(?!["\\/bfnrt])', r'\\\\', json_str)
        
        try:
            return json.loads(json_str)
        except json.JSONDecodeError as e:
            st.error(f"Erreur de parsing JSON: {e}")
            
            try:
                json_str_corrige = json_str.replace("\\", "\\\\")
                json_str_corrige = json_str_corrige.replace('\\"', '"')
                json_str_corrige = json_str_corrige.replace("\\\\n", "\\n")
                json_str_corrige = json_str_corrige.replace("\\\\'", "'")
                
                return json.loads(json_str_corrige)
            except json.JSONDecodeError:
                st.warning("Tentative d'extraction manuelle des données...")
                return extraire_donnees_manuellement(json_str)
    else:
        try:
            pattern2 = r'({[^{}]*"nom_consultant"[^}]*})'
            match2 = re.search(pattern2, str(reponse_ia), re.DOTALL)
            if match2:
                return json.loads(match2.group(1))
        except:
            pass
        
        st.error("Aucun contenu JSON trouvé dans la réponse")
        return None

def extraire_donnees_manuellement(json_str):
    """
    Extraction manuelle des données en cas d'échec du parsing JSON
    """
    donnees = {}
    champs = [
        'nom_consultant', 'titre_du_poste', 'niveaux_intervention',
        'hobbies_et_divers', 'connaissances', 'experiences', 
        'mois_debut_experience', 'nom_entreprise', 'points_forts', 'formations'
    ]
    
    for champ in champs:
        pattern = rf'"{champ}":\s*"([^"]*(?:\\.[^"]*)*)"'
        match = re.search(pattern, json_str, re.DOTALL)
        if match:
            valeur = match.group(1)
            valeur = valeur.replace('\\"', '"').replace('\\n', '\n').replace("\\'", "'")
            donnees[champ] = valeur
    
    return donnees if donnees else None

def parser_connaissances(connaissances_str):
    """
    Parse la chaîne de connaissances pour créer un dictionnaire
    """
    connaissances = {}
    
    if not connaissances_str:
        return connaissances
    
    # Nettoyer la chaîne
    connaissances_str = connaissances_str.replace('\\n', '\n')
    
    # Diviser par lignes
    lignes = connaissances_str.split('\n')
    
    for ligne in lignes:
        ligne = ligne.strip()
        if ':' in ligne and ligne:
            # Diviser par le premier ':'
            parties = ligne.split(':', 1)
            if len(parties) == 2:
                categorie = parties[0].strip()
                contenu = parties[1].strip()
                
                # Nettoyer les noms de catégories courants
                categorie_mappings = {
                    'Langages et Frameworks': 'Langages et Framework',
                    'Systèmes d\'exploitation': 'Systèmes d\'exploitation',
                    'DBMS/Servers': 'SGBD',
                    'DBMSServers': 'SGBD',
                    'Méthode de travail agile': 'Méthodologie',
                    'Outils d\'automatisations de taches (DEVOPS)': 'DevOps et Cloud',
                    'Outils d\'automatisations de taches': 'DevOps et Cloud'
                }
                
                categorie = categorie_mappings.get(categorie, categorie)
                connaissances[categorie] = contenu
    
    return connaissances

def creer_tableau_connaissances_a_position(doc, paragraph_position, connaissances_dict):
    """
    Crée un tableau de connaissances à une position spécifique dans le document
    """
    if not connaissances_dict:
        # Tableau par défaut si pas de données
        connaissances_dict = {
            'Langages et Framework': '.NET (C#, ASP.NET), MVC, WEB API, ANGULAR, TYPESCRIPT',
            'SGBD': 'MYSQL, POSTGRESQL, MONGODB, SQL Serveur',
            'Systèmes d\'exploitation': 'Linux (Ubuntu), Windows',
            'Outils': 'VsCode, GIT, GitHub, Visual studio',
            'DevOps et Cloud': 'DOCKER, KUBERNETES, CI/CD Devops',
            'Méthodologie': 'Agile SCRUM'
        }
    
    # Obtenir l'élément parent et l'index du paragraphe
    parent_element = paragraph_position._element.getparent()
    paragraph_index = list(parent_element).index(paragraph_position._element)
    
    # Créer l'élément table XML
    table_element = parent_element.makeelement(qn('w:tbl'))
    
    # Créer les propriétés de la table
    tbl_props = OxmlElement('w:tblPr')
    tbl_style = OxmlElement('w:tblStyle')
    tbl_style.set(qn('w:val'), 'TableGrid')
    tbl_props.append(tbl_style)
    
    tbl_width = OxmlElement('w:tblW')
    tbl_width.set(qn('w:w'), '5000')
    tbl_width.set(qn('w:type'), 'pct')
    tbl_props.append(tbl_width)
    
    table_element.append(tbl_props)
    
    # Couleur bleu Clinkast pour les en-têtes
    couleur_clinkast = "1F4E79"
    
    # Créer les lignes du tableau
    for categorie, contenu in connaissances_dict.items():
        # Créer une ligne
        tr_element = OxmlElement('w:tr')
        
        # Cellule 1 - Catégorie (fond blanc, texte bleu)
        tc1_element = OxmlElement('w:tc')
        tc1_props = OxmlElement('w:tcPr')
        
        # Largeur de la cellule
        tc1_width = OxmlElement('w:tcW')
        tc1_width.set(qn('w:w'), '2000')
        tc1_width.set(qn('w:type'), 'pct')
        tc1_props.append(tc1_width)
        
        # Couleur de fond blanche (pas de couleur de fond spécifique = blanc par défaut)
        # Pas de shd_element pour garder le fond blanc
        
        # Alignement vertical
        vAlign1 = OxmlElement('w:vAlign')
        vAlign1.set(qn('w:val'), 'center')
        tc1_props.append(vAlign1)
        
        tc1_element.append(tc1_props)
        
        # Paragraphe dans la cellule 1
        p1_element = OxmlElement('w:p')
        p1_props = OxmlElement('w:pPr')
        p1_align = OxmlElement('w:jc')
        p1_align.set(qn('w:val'), 'left')  # Alignement à gauche
        p1_props.append(p1_align)
        p1_element.append(p1_props)
        
        # Run avec le texte
        r1_element = OxmlElement('w:r')
        r1_props = OxmlElement('w:rPr')
        r1_bold = OxmlElement('w:b')
        r1_color = OxmlElement('w:color')
        r1_color.set(qn('w:val'), couleur_clinkast)  # Texte bleu (même couleur que l'ancien fond)
        r1_size = OxmlElement('w:sz')
        r1_size.set(qn('w:val'), '22')  # 11pt
        r1_props.append(r1_bold)
        r1_props.append(r1_color)
        r1_props.append(r1_size)
        r1_element.append(r1_props)
        
        r1_text = OxmlElement('w:t')
        r1_text.text = categorie
        r1_element.append(r1_text)
        
        p1_element.append(r1_element)
        tc1_element.append(p1_element)
        
        # Cellule 2 - Contenu
        tc2_element = OxmlElement('w:tc')
        tc2_props = OxmlElement('w:tcPr')
        
        # Largeur de la cellule
        tc2_width = OxmlElement('w:tcW')
        tc2_width.set(qn('w:w'), '3000')
        tc2_width.set(qn('w:type'), 'pct')
        tc2_props.append(tc2_width)
        
        tc2_element.append(tc2_props)
        
        # Paragraphe dans la cellule 2
        p2_element = OxmlElement('w:p')
        p2_props = OxmlElement('w:pPr')
        p2_align = OxmlElement('w:jc')
        p2_align.set(qn('w:val'), 'left')  # Alignement à gauche
        p2_props.append(p2_align)
        p2_element.append(p2_props)
        
        r2_element = OxmlElement('w:r')
        r2_props = OxmlElement('w:rPr')
        r2_size = OxmlElement('w:sz')
        r2_size.set(qn('w:val'), '18')  # 9pt
        r2_props.append(r2_size)
        r2_element.append(r2_props)
        
        r2_text = OxmlElement('w:t')
        r2_text.text = contenu
        r2_element.append(r2_text)
        
        p2_element.append(r2_element)
        tc2_element.append(p2_element)
        
        # Ajouter les cellules à la ligne
        tr_element.append(tc1_element)
        tr_element.append(tc2_element)
        
        # Ajouter la ligne au tableau
        table_element.append(tr_element)
    
    # Insérer le tableau après le paragraphe des connaissances
    parent_element.insert(paragraph_index + 1, table_element)
    
def parser_formation(formation_str):
    """
    Parse la chaîne de formation pour extraire année et intitulé
    """
    formations = {}
    
    if not formation_str:
        return formations
    
    # Nettoyer la chaîne
    formation_str = formation_str.replace('\\n', '\n')
    
    # Essayer de détecter les formats courants
    # Format 1: "YYYY Titre de formation"
    # Format 2: "Titre de formation YYYY"
    # Format 3: "Titre de formation (YYYY)"
    
    lignes = formation_str.split('\n') if '\n' in formation_str else [formation_str]
    
    for ligne in lignes:
        ligne = ligne.strip()
        if not ligne:
            continue
            
        # Chercher une année (4 chiffres)
        import re
        annee_match = re.search(r'\b(19|20)\d{2}\b', ligne)
        
        if annee_match:
            annee = annee_match.group()
            # Enlever l'année pour garder le titre
            titre = ligne.replace(annee, '').strip()
            # Nettoyer les caractères de ponctuation en trop
            titre = titre.strip('()- ,')
            formations[annee] = titre
        else:
            # Si pas d'année trouvée, utiliser "N/A" comme année
            formations['N/A'] = ligne
    
    return formations

def creer_tableau_formation_a_position(doc, paragraph_position, formation_str):
    """
    Crée un tableau de formation à une position spécifique
    """
    formations_dict = parser_formation(formation_str)
    
    if not formations_dict:
        # Formation par défaut
        formations_dict = {'2024': 'Formation à définir'}
    
    # Obtenir l'élément parent et l'index du paragraphe
    parent_element = paragraph_position._element.getparent()
    paragraph_index = list(parent_element).index(paragraph_position._element)
    
    # Créer l'élément table XML
    table_element = parent_element.makeelement(qn('w:tbl'))
    
    # Créer les propriétés de la table
    tbl_props = OxmlElement('w:tblPr')
    tbl_style = OxmlElement('w:tblStyle')
    tbl_style.set(qn('w:val'), 'TableGrid')
    tbl_props.append(tbl_style)
    
    tbl_width = OxmlElement('w:tblW')
    tbl_width.set(qn('w:w'), '5000')
    tbl_width.set(qn('w:type'), 'pct')
    tbl_props.append(tbl_width)
    
    table_element.append(tbl_props)
    
    # Couleur bleu Clinkast pour le texte
    couleur_clinkast = "1F4E79"
    
    # Créer les lignes du tableau
    for annee, formation in formations_dict.items():
        # Créer une ligne
        tr_element = OxmlElement('w:tr')
        
        # Cellule 1 - Année (fond blanc, texte bleu)
        tc1_element = OxmlElement('w:tc')
        tc1_props = OxmlElement('w:tcPr')
        
        # Largeur de la cellule
        tc1_width = OxmlElement('w:tcW')
        tc1_width.set(qn('w:w'), '1500')
        tc1_width.set(qn('w:type'), 'pct')
        tc1_props.append(tc1_width)
        
        tc1_element.append(tc1_props)
        
        # Paragraphe dans la cellule 1
        p1_element = OxmlElement('w:p')
        p1_props = OxmlElement('w:pPr')
        p1_align = OxmlElement('w:jc')
        p1_align.set(qn('w:val'), 'left')
        p1_props.append(p1_align)
        p1_element.append(p1_props)
        
        # Run avec le texte
        r1_element = OxmlElement('w:r')
        r1_props = OxmlElement('w:rPr')
        r1_bold = OxmlElement('w:b')
        r1_color = OxmlElement('w:color')
        r1_color.set(qn('w:val'), couleur_clinkast)
        r1_size = OxmlElement('w:sz')
        r1_size.set(qn('w:val'), '20')  # 10pt
        r1_props.append(r1_bold)
        r1_props.append(r1_color)
        r1_props.append(r1_size)
        r1_element.append(r1_props)
        
        r1_text = OxmlElement('w:t')
        r1_text.text = annee
        r1_element.append(r1_text)
        
        p1_element.append(r1_element)
        tc1_element.append(p1_element)
        
        # Cellule 2 - Formation
        tc2_element = OxmlElement('w:tc')
        tc2_props = OxmlElement('w:tcPr')
        
        # Largeur de la cellule
        tc2_width = OxmlElement('w:tcW')
        tc2_width.set(qn('w:w'), '3500')
        tc2_width.set(qn('w:type'), 'pct')
        tc2_props.append(tc2_width)
        
        tc2_element.append(tc2_props)
        
        # Paragraphe dans la cellule 2
        p2_element = OxmlElement('w:p')
        p2_props = OxmlElement('w:pPr')
        p2_align = OxmlElement('w:jc')
        p2_align.set(qn('w:val'), 'left')
        p2_props.append(p2_align)
        p2_element.append(p2_props)
        
        r2_element = OxmlElement('w:r')
        r2_props = OxmlElement('w:rPr')
        r2_size = OxmlElement('w:sz')
        r2_size.set(qn('w:val'), '18')  # 9pt
        r2_props.append(r2_size)
        r2_element.append(r2_props)
        
        r2_text = OxmlElement('w:t')
        r2_text.text = formation
        r2_element.append(r2_text)
        
        p2_element.append(r2_element)
        tc2_element.append(p2_element)
        
        # Ajouter les cellules à la ligne
        tr_element.append(tc1_element)
        tr_element.append(tc2_element)
        
        # Ajouter la ligne au tableau
        table_element.append(tr_element)
    
    # Insérer le tableau après le paragraphe
    parent_element.insert(paragraph_index + 1, table_element)
    
    return True

def creer_tableau_hobbies_a_position(doc, paragraph_position, hobbies_str):
    """
    Crée un tableau hobbies & divers à une position spécifique
    """
    # Structure fixe pour hobbies
    hobbies_dict = {
        'Langues': 'Français, Anglais (intermédiaire)',
        'Hobbies': hobbies_str if hobbies_str else 'À définir'
    }
    
    # Obtenir l'élément parent et l'index du paragraphe
    parent_element = paragraph_position._element.getparent()
    paragraph_index = list(parent_element).index(paragraph_position._element)
    
    # Créer l'élément table XML
    table_element = parent_element.makeelement(qn('w:tbl'))
    
    # Créer les propriétés de la table
    tbl_props = OxmlElement('w:tblPr')
    tbl_style = OxmlElement('w:tblStyle')
    tbl_style.set(qn('w:val'), 'TableGrid')
    tbl_props.append(tbl_style)
    
    tbl_width = OxmlElement('w:tblW')
    tbl_width.set(qn('w:w'), '5000')
    tbl_width.set(qn('w:type'), 'pct')
    tbl_props.append(tbl_width)
    
    table_element.append(tbl_props)
    
    # Couleur bleu Clinkast pour le texte
    couleur_clinkast = "1F4E79"
    
    # Créer les lignes du tableau
    for categorie, contenu in hobbies_dict.items():
        # Créer une ligne
        tr_element = OxmlElement('w:tr')
        
        # Cellule 1 - Catégorie (Langues/Hobbies)
        tc1_element = OxmlElement('w:tc')
        tc1_props = OxmlElement('w:tcPr')
        
        # Largeur de la cellule
        tc1_width = OxmlElement('w:tcW')
        tc1_width.set(qn('w:w'), '1500')
        tc1_width.set(qn('w:type'), 'pct')
        tc1_props.append(tc1_width)
        
        tc1_element.append(tc1_props)
        
        # Paragraphe dans la cellule 1
        p1_element = OxmlElement('w:p')
        p1_props = OxmlElement('w:pPr')
        p1_align = OxmlElement('w:jc')
        p1_align.set(qn('w:val'), 'left')
        p1_props.append(p1_align)
        p1_element.append(p1_props)
        
        # Run avec le texte
        r1_element = OxmlElement('w:r')
        r1_props = OxmlElement('w:rPr')
        r1_bold = OxmlElement('w:b')
        r1_color = OxmlElement('w:color')
        r1_color.set(qn('w:val'), couleur_clinkast)
        r1_size = OxmlElement('w:sz')
        r1_size.set(qn('w:val'), '20')  # 10pt
        r1_props.append(r1_bold)
        r1_props.append(r1_color)
        r1_props.append(r1_size)
        r1_element.append(r1_props)
        
        r1_text = OxmlElement('w:t')
        r1_text.text = categorie + ' :'
        r1_element.append(r1_text)
        
        p1_element.append(r1_element)
        tc1_element.append(p1_element)
        
        # Cellule 2 - Contenu
        tc2_element = OxmlElement('w:tc')
        tc2_props = OxmlElement('w:tcPr')
        
        # Largeur de la cellule
        tc2_width = OxmlElement('w:tcW')
        tc2_width.set(qn('w:w'), '3500')
        tc2_width.set(qn('w:type'), 'pct')
        tc2_props.append(tc2_width)
        
        tc2_element.append(tc2_props)
        
        # Paragraphe dans la cellule 2
        p2_element = OxmlElement('w:p')
        p2_props = OxmlElement('w:pPr')
        p2_align = OxmlElement('w:jc')
        p2_align.set(qn('w:val'), 'left')
        p2_props.append(p2_align)
        p2_element.append(p2_props)
        
        r2_element = OxmlElement('w:r')
        r2_props = OxmlElement('w:rPr')
        r2_size = OxmlElement('w:sz')
        r2_size.set(qn('w:val'), '18')  # 9pt
        r2_props.append(r2_size)
        r2_element.append(r2_props)
        
        r2_text = OxmlElement('w:t')
        r2_text.text = contenu
        r2_element.append(r2_text)
        
        p2_element.append(r2_element)
        tc2_element.append(p2_element)
        
        # Ajouter les cellules à la ligne
        tr_element.append(tc1_element)
        tr_element.append(tc2_element)
        
        # Ajouter la ligne au tableau
        table_element.append(tr_element)
    
    # Insérer le tableau après le paragraphe
    parent_element.insert(paragraph_index + 1, table_element)
    
    return True

def remplacer_texte_paragraph(paragraph, ancien_texte, nouveau_texte):
    """
    Remplace du texte dans un paragraphe en préservant le formatage
    """
    if ancien_texte in paragraph.text:
        # Parcourir tous les runs du paragraphe
        for run in paragraph.runs:
            if ancien_texte in run.text:
                run.text = run.text.replace(ancien_texte, nouveau_texte)
                return True
        
        # Si le placeholder est réparti sur plusieurs runs
        texte_complet = paragraph.text
        if ancien_texte in texte_complet:
            # Effacer tout le contenu du paragraphe
            paragraph.clear()
            # Ajouter le nouveau texte
            nouveau_run = paragraph.add_run(texte_complet.replace(ancien_texte, nouveau_texte))
            return True
    
    return False

def ajouter_liste_paragraphes(doc, paragraph_parent, items, style_bullet=True):
    """
    Ajoute une série de paragraphes formatés en liste
    """
    parent_element = paragraph_parent._element.getparent()
    
    for item in items:
        # Créer un nouveau paragraphe
        new_p_element = parent_element.makeelement(qn('w:p'))
        parent_element.insert(parent_element.index(paragraph_parent._element) + 1, new_p_element)
        
        # Créer le paragraphe Python-docx
        new_paragraph = paragraph_parent.__class__(new_p_element, paragraph_parent._parent)
        
        if style_bullet:
            # Ajouter la puce
            run = new_paragraph.add_run(f"• {item.strip()}")
        else:
            run = new_paragraph.add_run(item.strip())
        
        run.font.size = Pt(10)

def remplacer_placeholders(doc, data):
    """
    Remplace tous les placeholders dans le document avec les données
    """
    # Dictionnaire des remplacements
    remplacements = {
        '{{nom_consultant}}': data.get('nom_consultant', 'Nom du consultant'),
        '{{titre_du_poste}}': data.get('titre_du_poste', 'Titre du poste'),
        '{{points_forts}}': data.get('points_forts', 'Points forts à définir'),
        '{{niveaux_intervention}}': data.get('niveaux_intervention', 'Niveaux d\'intervention à définir'),
        '{{formation}}': data.get('formations', 'Formation à définir'),
        '{{hobbies_et_divers}}': data.get('hobbies_et_divers', 'Hobbies et divers à définir'),
        '{{experiences}}': data.get('experiences', 'Expériences à définir'),
        '{{mois_debut_experience}}': data.get('mois_debut_experience', 'Date'),
        '{{nom_entreprise}}': data.get('nom_entreprise', 'Entreprise')
    }
    
    # Initialiser les variables de contrôle pour tous les tableaux
    tableau_connaissances_cree = False
    tableau_formation_cree = False  
    tableau_hobbies_cree = False
    
    # Traiter le tableau des connaissances d'abord
    tableau_cree = False
    paragraphs_to_process = list(doc.paragraphs)  # Créer une copie de la liste
    
    for i, paragraph in enumerate(paragraphs_to_process):
        if '{{tableau_connaissances}}' in paragraph.text and not tableau_cree:
            # Parser les connaissances
            connaissances_str = data.get('connaissances', '')
            connaissances_dict = parser_connaissances(connaissances_str)
            
            # Effacer le contenu du paragraphe
            paragraph.clear()
            
            # Créer le tableau à cette position exacte
            creer_tableau_connaissances_a_position(doc, paragraph, connaissances_dict)
            tableau_cree = True
            continue
    
    # Remplacer dans tous les paragraphes (y compris les placeholders de tableaux non traités)
    for paragraph in doc.paragraphs:
        # Traiter les placeholders de tableaux qui n'ont pas été trouvés dans la première passe
        if '{{tableau_formation}}' in paragraph.text and not tableau_formation_cree:
            formation_str = data.get('formations', '')
            remplacer_texte_paragraph(paragraph, '{{tableau_formation}}', '')
            creer_tableau_formation_a_position(doc, paragraph, formation_str)
            tableau_formation_cree = True
            continue
        
        if '{{tableau_hobbies}}' in paragraph.text and not tableau_hobbies_cree:
            hobbies_str = data.get('hobbies_et_divers', '')
            remplacer_texte_paragraph(paragraph, '{{tableau_hobbies}}', '')
            creer_tableau_hobbies_a_position(doc, paragraph, hobbies_str)
            tableau_hobbies_cree = True
            continue
        
        if '{{tableau_connaissances}}' in paragraph.text and not tableau_connaissances_cree:
            connaissances_str = data.get('connaissances', '')
            connaissances_dict = parser_connaissances(connaissances_str)
            remplacer_texte_paragraph(paragraph, '{{tableau_connaissances}}', '')
            creer_tableau_connaissances_a_position(doc, paragraph, connaissances_dict)
            tableau_connaissances_cree = True
            continue
        
        # Remplacements normaux
        for placeholder, valeur in remplacements.items():
            if placeholder in paragraph.text:
                # Traitement spécial pour les listes (points forts, niveaux d'intervention)
                if placeholder in ['{{points_forts}}', '{{niveaux_intervention}}'] and ',' in valeur:
                    items = [item.strip() for item in valeur.split(',')]
                    if items:
                        # Remplacer le placeholder par le premier item
                        remplacer_texte_paragraph(paragraph, placeholder, f"• {items[0]}")
                        
                        # Ajouter les autres items
                        if len(items) > 1:
                            ajouter_liste_paragraphes(doc, paragraph, items[1:], style_bullet=True)
                else:
                    # Remplacement simple
                    remplacer_texte_paragraph(paragraph, placeholder, valeur)
    
    # Remplacer dans les tableaux existants
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, valeur in remplacements.items():
                        if placeholder in paragraph.text:
                            remplacer_texte_paragraph(paragraph, placeholder, valeur)

def generer_cv_depuis_template(template_file, data):
    """
    Génère un CV en remplissant un template Word
    """
    try:
        # Charger le template
        doc = Document(template_file)
        
        # Remplacer tous les placeholders
        remplacer_placeholders(doc, data)
        
        return doc
        
    except Exception as e:
        st.error(f"Erreur lors du traitement du template: {str(e)}")
        return None

def main():
    st.title("📄 Générateur de CV Clinkast - Template Personnalisé")
    st.markdown("*Utilisez votre propre template Word avec placeholders*")
    st.markdown("---")
    
    # Section template
    st.subheader("📂 1. Chargez votre template Word")
    
    col_template1, col_template2 = st.columns([2, 1])
    
    with col_template1:
        template_file = st.file_uploader(
            "Sélectionnez votre template Word (.docx)",
            type=['docx'],
            help="Fichier Word contenant les placeholders comme {{nom_consultant}}, {{tableau_connaissances}}, etc."
        )
    
    with col_template2:
        st.info("""
        **Placeholders supportés:**
        - `{{nom_consultant}}`
        - `{{titre_du_poste}}`
        - `{{points_forts}}`
        - `{{niveaux_intervention}}`
        - `{{tableau_formation}}`
        - `{{tableau_connaissances}}`
        - `{{tableau_hobbies}}`
        - `{{experiences}}`
        """)
    
    # Section données IA
    st.subheader("🤖 2. Données de l'IA")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        # Zone de texte pour la réponse IA
        reponse_ia = st.text_area(
            "Réponse IA",
            height=200,
            placeholder="Collez ici la réponse complète de votre IA (ChatCompletion...)",
            help="Collez la réponse complète incluant ChatCompletion(...)"
        )
        
        # Exemple de données pour test
        if st.checkbox("Utiliser un exemple de données pour tester"):
            exemple_reponse = """ChatCompletion(id='chatcmpl-C6HYXzuV6UnjUGX3XwucQiF8qY3iR', choices=[Choice(finish_reason='stop', index=0, logprobs=None, message=ChatCompletionMessage(content='{\n    "niveaux_intervention": "Développeur Full Stack, Chef d\'équipe, Devops",\n    "hobbies_et_divers": "Musique, danse, Nouvelles technologies Informatiques",\n    "connaissances": "Langages et Frameworks: HTML, CSS, SQL, C#, Java, Python, JavaScript, Yaml, TypeScript, .Net, ASP.Net, Entity Framework, Razor, Angular, Angular Material, Azure storage, React Js, Bootstrap, Tailwind CSS, Spring Boot, Flutter\\nSystèmes d\'exploitation: Windows, Linux\\nOutils: Visual Studio, Visual Studio Code, Balsamiq Cloud, Eclipse, Anaconda, Git, Azure Devops\\nDBMS/Servers: MySQL, SQL Server 2019, PostgreSQL\\nMéthode de travail agile: SCRUM\\nOutils d\'automatisations de taches (DEVOPS): Ansible, Dockers, Kubernetes, Flux CD",\n    "titre_du_poste": "Développeur C# / .NET",\n    "experiences": "Développeur Full Stack And Software AI engineer chez Clinkast France, Développeur Full Stack chez LuxSoft France, Développeur Full Stack + Devops + Chef d'équipe chez Cethia SARL",\n    "mois_debut_experience": "Novembre 2024",\n    "nom_entreprise": "Clinkast France",\n    "nom_consultant": "FEUZING NTEMMA Donald",\n    "points_forts": "Expérience en développement Full Stack, Connaissance en C#, .NET, Connaissance en Devops, Expérience en gestion d\'équipe, Capacité d\'adaptation, Créativité",\n    "formations": "Master en Intelligence Artificielle (Major de Promotion avec Mention TB) Juillet 2024, Licence En Informatique Juin 2022"\n}', refusal=None, role='assistant', annotations=[], audio=None, function_call=None, tool_calls=None))], created=1755613489, model='gpt-4-0613', object='chat.completion', service_tier='default', system_fingerprint=None, usage=CompletionUsage(completion_tokens=409, prompt_tokens=3470, total_tokens=3879, completion_tokens_details=CompletionTokensDetails(accepted_prediction_tokens=0, audio_tokens=0, reasoning_tokens=0, rejected_prediction_tokens=0), prompt_tokens_details=PromptTokensDetails(audio_tokens=0, cached_tokens=0)))"""
            reponse_ia = exemple_reponse
    
    with col2:
        st.subheader("⚙️ Options")
        nom_fichier = st.text_input("Nom du fichier CV", value="CV_genere_template.docx")
        
        # Bouton pour générer le CV
        if st.button("🚀 Générer le CV", type="primary", disabled=not template_file):
            if template_file and reponse_ia.strip():
                with st.spinner("Génération du CV depuis le template..."):
                    # Extraire les données
                    data = extraire_contenu_json(reponse_ia)
                    
                    if data:
                        st.success("✅ Données extraites avec succès !")
                        
                        # Afficher les données extraites
                        with st.expander("📊 Données extraites", expanded=False):
                            st.json(data)
                        
                        # Afficher les connaissances parsées
                        connaissances_dict = parser_connaissances(data.get('connaissances', ''))
                        if connaissances_dict:
                            with st.expander("📋 Tableau des connaissances", expanded=False):
                                for cat, cont in connaissances_dict.items():
                                    st.write(f"**{cat}:** {cont}")
                        
                        try:
                            # Générer le document Word depuis le template
                            doc = generer_cv_depuis_template(template_file, data)
                            
                            if doc:
                                # Sauvegarder temporairement
                                doc.save(nom_fichier)
                                
                                st.success(f"🎉 CV généré avec succès : {nom_fichier}")
                                
                                # Proposer le téléchargement
                                with open(nom_fichier, "rb") as file:
                                    btn = st.download_button(
                                        label="📥 Télécharger le CV",
                                        data=file,
                                        file_name=nom_fichier,
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                    )
                            
                        except Exception as e:
                            st.error(f"❌ Erreur lors de la génération du CV : {str(e)}")
                            st.exception(e)
                    else:
                        st.error("❌ Impossible d'extraire les données de la réponse IA")
            elif not template_file:
                st.warning("⚠️ Veuillez charger un template Word")
            else:
                st.warning("⚠️ Veuillez coller la réponse de l'IA")
    
    # Guide d'utilisation
    st.markdown("---")
    st.subheader("📚 Guide d'utilisation")
    
    col_guide1, col_guide2 = st.columns(2)
    
    with col_guide1:
        st.markdown("""
        **Étapes :**
        1. 📁 **Chargez votre template Word** contenant les placeholders
        2. 🤖 **Collez la réponse de l'IA** avec les données JSON
        3. ⚙️ **Configurez le nom du fichier** de sortie
        4. 🚀 **Générez le CV** - le template sera rempli automatiquement
        5. 📥 **Téléchargez** votre CV personnalisé
        """)
    
    with col_guide2:
        st.markdown("""
        **Format du template :**
        - Utilisez des placeholders comme `{{nom_consultant}}`
        - Pour le tableau des connaissances : `{{tableau_connaissances}}`
        - Les données de l'IA doivent contenir les champs correspondants
        - Le tableau sera créé automatiquement avec formatage Clinkast
        - Tous les autres éléments du template sont préservés
        """)
    
    # Exemple de template
    with st.expander("📄 Exemple de template Word", expanded=False):
        st.code("""
Template Word contenant :

CLINKAST                    Dossier de compétence

{{nom_consultant}} -- {{titre_du_poste}}

Points forts
{{points_forts}}

Niveaux d'Intervention  
{{niveaux_intervention}}

Formation
{{tableau_formation}}

Connaissances
{{tableau_connaissances}}

Hobbies & Divers
{{tableau_hobbies}}

Expériences Professionnelles
{{experiences}}

[En-tête et pied de page avec logo Clinkast]
        """)

if __name__ == "__main__":
    main()