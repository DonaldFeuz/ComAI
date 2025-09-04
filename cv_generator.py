import streamlit as st
import json
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement, qn
import io
import base64

# Configuration de la page
st.set_page_config(
    page_title="Générateur de CV",
    page_icon="📄",
    layout="wide"
)

def extraire_contenu_json(reponse_ia):
    """
    Extrait le contenu JSON de la réponse de l'IA
    """
    # Recherche du pattern JSON dans la chaîne de caractères
    pattern = r'content=\'({.*?})\''
    match = re.search(pattern, str(reponse_ia), re.DOTALL)
    
    if match:
        json_str = match.group(1)
        
        # Nettoyage approfondi des échappements
        # D'abord, remplacer les échappements doubles
        json_str = json_str.replace('\\\\', '\\')
        
        # Remplacer les apostrophes échappées
        json_str = json_str.replace("\\'", "'")
        
        # Traitement spécial pour les \n dans les chaînes JSON
        # Rechercher et remplacer les \\n par des \n réels
        json_str = json_str.replace('\\n', '\n')
        
        # Corriger les échappements invalides dans les chaînes JSON
        # Remplacer les \ non suivis d'un caractère d'échappement valide
        json_str = re.sub(r'\\(?!["\\/bfnrt])', r'\\\\', json_str)
        
        try:
            return json.loads(json_str)
        except json.JSONDecodeError as e:
            st.error(f"Erreur de parsing JSON: {e}")
            
            # Tentative de correction supplémentaire
            try:
                # Essayer de corriger les échappements problématiques
                json_str_corrige = json_str.replace("\\", "\\\\")
                json_str_corrige = json_str_corrige.replace('\\"', '"')
                json_str_corrige = json_str_corrige.replace("\\\\n", "\\n")
                json_str_corrige = json_str_corrige.replace("\\\\'", "'")
                
                return json.loads(json_str_corrige)
            except json.JSONDecodeError:
                # Dernier recours : extraction manuelle
                st.warning("Tentative d'extraction manuelle des données...")
                return extraire_donnees_manuellement(json_str)
    else:
        # Essayer de trouver juste le JSON brut
        try:
            # Recherche d'un objet JSON direct
            pattern2 = r'({[^{}]*"nom_consultant"[^}]*})'
            match2 = re.search(pattern2, str(reponse_ia), re.DOTALL)
            if match2:
                return json.loads(match2.group(1))
        except:
            pass
        
        st.error("Aucun contenu JSON trouvé dans la réponse")
        return None

def extraire_donnees_manuellement(json_str):
    """
    Extraction manuelle des données en cas d'échec du parsing JSON
    """
    donnees = {}
    
    # Dictionnaire des champs à extraire
    champs = [
        'nom_consultant', 'titre_du_poste', 'niveaux_intervention',
        'hobbies_et_divers', 'connaissances', 'experiences', 
        'mois_debut_experience', 'nom_entreprise', 'points_forts', 'formations'
    ]
    
    for champ in champs:
        # Pattern pour extraire chaque champ
        pattern = rf'"{champ}":\s*"([^"]*(?:\\.[^"]*)*)"'
        match = re.search(pattern, json_str, re.DOTALL)
        if match:
            valeur = match.group(1)
            # Nettoyer les échappements
            valeur = valeur.replace('\\"', '"').replace('\\n', '\n').replace("\\'", "'")
            donnees[champ] = valeur
    
    return donnees if donnees else None

def creer_en_tete_tableau(cell, text, color_rgb):
    """
    Crée un en-tête de tableau avec fond coloré
    """
    # Ajouter le texte
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(255, 255, 255)  # Texte blanc
    
    # Définir la couleur de fond
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), f"{color_rgb:06x}")
    cell._tc.get_or_add_tcPr().append(shading_elm)
    
    # Centrer le texte
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def generer_cv_word(data):
    """
    Génère un document Word CV à partir des données extraites
    """
    # Créer un nouveau document
    doc = Document()
    
    # Configuration des marges
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Titre principal
    titre = doc.add_heading(data.get('nom_consultant', 'Consultant'), level=1)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Sous-titre
    sous_titre = doc.add_paragraph(data.get('titre_du_poste', ''))
    sous_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sous_titre.runs:
        sous_titre.runs[0].bold = True
    
    # Extraire les informations de connaissances
    connaissances_brutes = data.get('connaissances', '')
    
    # Parser les connaissances pour extraire les sections
    sections_connaissances = {}
    if connaissances_brutes:
        lignes = connaissances_brutes.split('\\n')
        for ligne in lignes:
            if ':' in ligne:
                cle, valeur = ligne.split(':', 1)
                sections_connaissances[cle.strip()] = valeur.strip()
    
    # Créer le tableau des connaissances
    sections_data = [
        ("Langages et Framework", sections_connaissances.get('Langages et Frameworks', 
         '.NET (C#, ASP.NET), MVC, WEB API, ANGULAR, TYPESCRIPT, ANGULAR MATERIAL, RXJS, ENTITY FRAMEWORK, LINQ, JAVASCRIPT, JSON, SOAP, SQL, RAZOR, YAML, BOOTSTRAP, TAILWIND'), 0x1F4E79),
        ("SGBD", "MYSQL, POSTGRESQL, MONGODB, SQL Serveur", 0x1F4E79),
        ("Systèmes d'exploitation", sections_connaissances.get('Systèmes d\'exploitation', 'Linux (Ubuntu), Windows'), 0x1F4E79),
        ("Outils", sections_connaissances.get('Outils', 'VsCode, GIT, GitHub, gitLAB, Visual studio, Ansible, Flux CD, Balsamiq Cloud, Eclipse, Anaconda, Azure Devops'), 0x1F4E79),
        ("DevOps et Cloud", sections_connaissances.get('Outils d\'automatisations de taches (DEVOPS)', 'DOCKER, KUBERNETES, CI/CD Devops'), 0x1F4E79),
        ("Méthodologie", sections_connaissances.get('Méthode de travail agile', 'Agile SCRUM'), 0x1F4E79),
        ("Autres", "Injection de dépendances, Test unitaire automatisé, Architecture en oignon, Clean, Architecture, Devops", 0x1F4E79)
    ]
    
    # Ajouter le titre Connaissances
    doc.add_paragraph()  # Espace
    doc.add_heading('Connaissances', level=2)
    
    # Créer le tableau
    table = doc.add_table(rows=len(sections_data), cols=2)
    table.style = 'Table Grid'
    
    for i, (titre, contenu, couleur) in enumerate(sections_data):
        # Cellule de titre
        creer_en_tete_tableau(table.cell(i, 0), titre, couleur)
        
        # Cellule de contenu
        cell_contenu = table.cell(i, 1)
        paragraph = cell_contenu.paragraphs[0]
        paragraph.clear()
        run = paragraph.add_run(contenu)
        run.font.size = Pt(10)
        
        # Ajuster la largeur des colonnes
        table.cell(i, 0).width = Inches(2.0)
        table.cell(i, 1).width = Inches(4.5)
    
    # Section Expériences
    if data.get('experiences'):
        doc.add_paragraph()
        doc.add_heading('Expériences Professionnelles', level=2)
        
        exp_paragraph = doc.add_paragraph()
        run = exp_paragraph.add_run(data.get('experiences', ''))
        run.font.size = Pt(10)
        
        # Ajouter les détails de l'expérience actuelle
        if data.get('nom_entreprise'):
            detail_exp = doc.add_paragraph()
            detail_run = detail_exp.add_run(f"Entreprise actuelle: {data.get('nom_entreprise')} (depuis {data.get('mois_debut_experience', 'N/A')})")
            detail_run.font.size = Pt(10)
            detail_run.bold = True
    
    # Section Formations
    if data.get('formations'):
        doc.add_paragraph()
        doc.add_heading('Formations', level=2)
        formation_paragraph = doc.add_paragraph()
        run = formation_paragraph.add_run(data.get('formations', ''))
        run.font.size = Pt(10)
    
    # Section Points forts
    if data.get('points_forts'):
        doc.add_paragraph()
        doc.add_heading('Points Forts', level=2)
        points_paragraph = doc.add_paragraph()
        run = points_paragraph.add_run(data.get('points_forts', ''))
        run.font.size = Pt(10)
    
    # Section Niveaux d'intervention
    if data.get('niveaux_intervention'):
        doc.add_paragraph()
        doc.add_heading('Niveaux d\'intervention', level=2)
        niveaux_paragraph = doc.add_paragraph()
        run = niveaux_paragraph.add_run(data.get('niveaux_intervention'))
        run.font.size = Pt(10)
    
    # Section Hobbies
    if data.get('hobbies_et_divers'):
        doc.add_paragraph()
        doc.add_heading('Hobbies et Divers', level=2)
        hobbies_paragraph = doc.add_paragraph()
        run = hobbies_paragraph.add_run(data.get('hobbies_et_divers'))
        run.font.size = Pt(10)
    
    return doc

def get_binary_file_downloader_html(bin_file, file_label='File'):
    """
    Générer un lien de téléchargement pour les fichiers binaires
    """
    with open(bin_file, 'rb') as f:
        data = f.read()
    bin_str = base64.b64encode(data).decode()
    href = f'<a href="data:application/octet-stream;base64,{bin_str}" download="{bin_file}">📥 Télécharger {file_label}</a>'
    return href

def main():
    st.title("📄 Générateur de CV à partir de réponse IA")
    st.markdown("---")
    
    # Interface utilisateur
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.subheader("Collez la réponse de votre IA ici :")
        
        # Zone de texte pour la réponse IA
        reponse_ia = st.text_area(
            "Réponse IA",
            height=200,
            placeholder="Collez ici la réponse complète de votre IA (ChatCompletion...)",
            help="Collez la réponse complète incluant ChatCompletion(...)"
        )
        
        # Exemple de données pour test
        if st.checkbox("Utiliser un exemple de données pour tester"):
            exemple_reponse = """ChatCompletion(id='chatcmpl-C6HYXzuV6UnjUGX3XwucQiF8qY3iR', choices=[Choice(finish_reason='stop', index=0, logprobs=None, message=ChatCompletionMessage(content='{\n    "niveaux_intervention": "Développeur Full Stack, Chef d\'équipe, Devops",\n    "hobbies_et_divers": "Musique, danse, Nouvelles technologies Informatiques",\n    "connaissances": "Langages et Frameworks: HTML, CSS, SQL, C#, Java, Python, JavaScript, Yaml, TypeScript, .Net, ASP.Net, Entity Framework, Razor, Angular, Angular Material, Azure storage, React Js, Bootstrap, Tailwind CSS, Spring Boot, Flutter\\nSystèmes d\'exploitation: Windows, Linux\\nOutils: Visual Studio, Visual Studio Code, Balsamiq Cloud, Eclipse, Anaconda, Git, Azure Devops\\nDBMS/Servers: MySQL, SQL Server 2019, PostgreSQL\\nMéthode de travail agile: SCRUM\\nOutils d\'automatisations de taches (DEVOPS): Ansible, Dockers, Kubernetes, Flux CD",\n    "titre_du_poste": "Développeur C# / .NET",\n    "experiences": "Développeur Full Stack And Software AI engineer chez Clinkast France, Développeur Full Stack chez LuxSoft France, Développeur Full Stack + Devops + Chef d'équipe chez Cethia SARL",\n    "mois_debut_experience": "Novembre 2024",\n    "nom_entreprise": "Clinkast France",\n    "nom_consultant": "FEUZING NTEMMA Donald",\n    "points_forts": "Expérience en développement Full Stack, Connaissance en C#, .NET, Connaissance en Devops, Expérience en gestion d\'équipe, Capacité d\'adaptation, Créativité",\n    "formations": "Master en Intelligence Artificielle (Major de Promotion avec Mention TB) Juillet 2024, Licence En Informatique Juin 2022"\n}', refusal=None, role='assistant', annotations=[], audio=None, function_call=None, tool_calls=None))], created=1755613489, model='gpt-4-0613', object='chat.completion', service_tier='default', system_fingerprint=None, usage=CompletionUsage(completion_tokens=409, prompt_tokens=3470, total_tokens=3879, completion_tokens_details=CompletionTokensDetails(accepted_prediction_tokens=0, audio_tokens=0, reasoning_tokens=0, rejected_prediction_tokens=0), prompt_tokens_details=PromptTokensDetails(audio_tokens=0, cached_tokens=0)))"""
            reponse_ia = exemple_reponse
    
    with col2:
        st.subheader("Options :")
        nom_fichier = st.text_input("Nom du fichier CV", value="CV_genere.docx")
        
        # Bouton pour générer le CV
        if st.button("🚀 Générer le CV", type="primary"):
            if reponse_ia.strip():
                with st.spinner("Génération du CV en cours..."):
                    # Extraire les données
                    data = extraire_contenu_json(reponse_ia)
                    
                    if data:
                        st.success("✅ Données extraites avec succès !")
                        
                        # Afficher les données extraites
                        with st.expander("📊 Données extraites", expanded=False):
                            st.json(data)
                        
                        try:
                            # Générer le document Word
                            doc = generer_cv_word(data)
                            
                            # Sauvegarder temporairement
                            doc.save(nom_fichier)
                            
                            st.success(f"🎉 CV généré avec succès : {nom_fichier}")
                            
                            # Proposer le téléchargement
                            with open(nom_fichier, "rb") as file:
                                btn = st.download_button(
                                    label="📥 Télécharger le CV",
                                    data=file,
                                    file_name=nom_fichier,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                                
                        except Exception as e:
                            st.error(f"❌ Erreur lors de la génération du CV : {str(e)}")
                    else:
                        st.error("❌ Impossible d'extraire les données de la réponse IA")
            else:
                st.warning("⚠️ Veuillez coller la réponse de l'IA")
    
    # Informations sur l'utilisation
    st.markdown("---")
    st.subheader("ℹ️ Comment utiliser cette application :")
    
    col_info1, col_info2 = st.columns(2)
    
    with col_info1:
        st.markdown("""
        **Étapes :**
        1. Copiez la réponse complète de votre IA (ChatCompletion...)
        2. Collez-la dans la zone de texte
        3. Donnez un nom à votre fichier CV
        4. Cliquez sur "Générer le CV"
        5. Téléchargez votre CV Word formaté
        """)
    
    with col_info2:
        st.markdown("""
        **Format attendu :**
        - La réponse doit contenir un JSON avec les champs :
          - `nom_consultant`
          - `titre_du_poste`
          - `connaissances`
          - `experiences`
          - `formations`
          - `points_forts`
          - etc.
        """)

if __name__ == "__main__":
    main()